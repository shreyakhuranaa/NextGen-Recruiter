import io
import json
import re

from docx import Document
from PyPDF2 import PdfReader


SKILL_BANK = [
    "python",
    "java",
    "javascript",
    "typescript",
    "react",
    "node",
    "flask",
    "fastapi",
    "sql",
    "postgresql",
    "mysql",
    "mongodb",
    "docker",
    "kubernetes",
    "aws",
    "gcp",
    "azure",
    "machine learning",
    "deep learning",
    "openai",
    "llm",
    "rag",
    "tailwind",
    "html",
    "css",
    "git",
]

SECTION_HINTS = {
    "projects": {"projects", "project", "academic projects", "personal projects"},
    "experience": {"experience", "work experience", "internship", "internships"},
    "education": {"education", "academics", "academic background"},
}


def parse_resume_file(file_storage):
    filename = (getattr(file_storage, "filename", "") or "").lower()
    raw = file_storage.read()
    text = _extract_text(raw, filename)
    parsed = _parse_text(text)
    return text, parsed


def _extract_text(raw: bytes, filename: str):
    if filename.endswith(".pdf"):
        reader = PdfReader(io.BytesIO(raw))
        return "\n".join((page.extract_text() or "") for page in reader.pages).strip()

    if filename.endswith(".docx"):
        doc = Document(io.BytesIO(raw))
        chunks = [paragraph.text for paragraph in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                chunks.append(" ".join(cell.text for cell in row.cells))
        return "\n".join(chunks).strip()

    raise ValueError("Only PDF and DOCX resumes are supported.")


def _parse_text(text: str):
    lines = [re.sub(r"\s+", " ", line).strip(" -|•\t") for line in text.splitlines()]
    lines = [line for line in lines if line]
    lower_text = text.lower()
    sections = _split_sections(lines)
    name = _guess_name(lines)
    email = _extract_email(text)
    phone = _extract_phone(text)
    skills = []
    for skill in SKILL_BANK:
        if re.search(rf"(?<![a-z0-9]){re.escape(skill)}(?![a-z0-9])", lower_text):
            skills.append(skill.title().replace("Aws", "AWS").replace("Gcp", "GCP").replace("Llm", "LLM"))

    projects = _top_items(sections.get("projects", []), 4)
    experience = _top_items(sections.get("experience", []), 4)
    education = _top_items(sections.get("education", []), 3)
    summary = _build_summary(name, skills, experience, projects)

    return {
        "name": name,
        "email": email,
        "phone": phone,
        "skills": skills[:16],
        "projects": projects,
        "experience": experience,
        "education": education,
        "summary": summary,
        "wordCount": len(re.findall(r"\b\w+\b", text)),
    }


def _split_sections(lines):
    sections = {"projects": [], "experience": [], "education": []}
    current = "experience"
    for line in lines:
        normalized = re.sub(r"[^a-zA-Z ]", "", line).strip().lower()
        matched = next((key for key, hints in SECTION_HINTS.items() if normalized in hints), None)
        if matched:
            current = matched
            continue
        sections.setdefault(current, []).append(line)
    return sections


def _guess_name(lines):
    for line in lines[:6]:
        if "@" in line or re.search(r"\d", line):
            continue
        if 1 < len(line.split()) <= 5:
            return line
    return ""


def _extract_email(text):
    match = re.search(r"[\w.+-]+@[\w-]+(?:\.[\w-]+)+", text)
    return match.group(0) if match else ""


def _extract_phone(text):
    for candidate in re.findall(r"\+?\d[\d\s().-]{8,}\d", text):
        digits = re.sub(r"\D", "", candidate)
        if 10 <= len(digits) <= 13:
            return candidate.strip()
    return ""


def _top_items(lines, limit):
    items = []
    for line in lines:
        if len(line) < 10:
            continue
        items.append(line[:220])
        if len(items) >= limit:
            break
    return items


def _build_summary(name, skills, experience, projects):
    skill_text = ", ".join(skills[:6]) or "general engineering skills"
    exp_text = experience[0] if experience else "limited prior experience provided"
    project_text = projects[0] if projects else "no project details detected"
    return (
        f"{name or 'Candidate'} appears strongest in {skill_text}. "
        f"Experience highlight: {exp_text}. Project highlight: {project_text}."
    )


def dump_parsed_resume(parsed):
    return json.dumps(parsed)
