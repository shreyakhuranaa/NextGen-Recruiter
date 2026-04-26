import os, io, time, logging, json, re, random, sqlite3, html, sys
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Optional, Tuple
os.environ.setdefault("TF_CPP_MIN_LOG_LEVEL", "3")
import streamlit as st
try:
    from openai import AzureOpenAI, OpenAI
    _OAI = True
except Exception:
    AzureOpenAI = None; OpenAI = None; _OAI = False

# streamlit-webrtc for real audio capture (report Ch.3 §4 — WebRTC integration)
try:
    from streamlit_webrtc import webrtc_streamer, WebRtcMode, RTCConfiguration
    import av
    _WEBRTC = True
except Exception:
    _WEBRTC = False

# Audio threading support for WebRTC processor
import queue, threading

log = logging.getLogger("ngr")
if not log.handlers:
    logging.basicConfig(level=logging.WARNING)
log.setLevel(logging.WARNING)


# ══════════════════════════════════════════════════════════════════════════════
# WEBRTC AUDIO PROCESSOR (report Ch.3 §6 — Audio Capture + VAD)
# Collects PCM frames from WebRTC stream into a global queue.
# VAD (Voice Activity Detection): stops collecting after 1.5s of silence.
# ══════════════════════════════════════════════════════════════════════════════
if _WEBRTC:
    import numpy as np
    from streamlit_webrtc import AudioProcessorBase

    class InterviewAudioProcessor(AudioProcessorBase):
        """
        Buffers raw PCM audio frames from WebRTC stream.
        Report Ch.3 §6: InterviewAudioProcessor with audio buffering + VAD logic.
        """
        def __init__(self):
            self._frames: list = []
            self._sample_rate: int = 48000
            self._lock = threading.Lock()

        def recv(self, frame: "av.AudioFrame") -> "av.AudioFrame":
            """Receive audio frame, convert to int16, store in buffer."""
            try:
                # Convert to numpy int16 PCM
                pcm = frame.to_ndarray()
                # av returns float32 fltp (planar) — convert to int16
                if pcm.dtype != np.int16:
                    pcm = np.clip(pcm * 32767, -32768, 32767).astype(np.int16)
                # Flatten multi-channel to mono
                if pcm.ndim > 1:
                    pcm = pcm.mean(axis=0).astype(np.int16)
                self._sample_rate = frame.sample_rate
                with self._lock:
                    self._frames.append(pcm)
            except Exception as e:
                pass
            return frame

        def get_wav(self) -> Optional[bytes]:
            """
            Assemble buffered frames into a WAV file with VAD trimming.
            Report Ch.3 §6: WAV conversion pipeline + silence detection.
            """
            import wave
            with self._lock:
                if not self._frames:
                    return None
                audio = np.concatenate(self._frames).flatten()

            if len(audio) < 2000:
                return None

            # VAD: find last active frame (RMS above threshold)
            chunk = max(int(self._sample_rate * 0.02), 160)  # 20ms chunks
            threshold = max(float(np.abs(audio).max()) * 0.01, 100.0)
            last_active = len(audio)
            for i in range(len(audio) - chunk, chunk, -chunk):
                if float(np.abs(audio[i:i+chunk]).mean()) > threshold:
                    last_active = min(i + chunk * 3, len(audio))
                    break

            audio = audio[:last_active]
            if len(audio) < 2000:
                return None

            buf = io.BytesIO()
            with wave.open(buf, 'wb') as wf:
                wf.setnchannels(1)
                wf.setsampwidth(2)
                wf.setframerate(self._sample_rate)
                wf.writeframes(audio.tobytes())
            return buf.getvalue()

        def clear(self):
            with self._lock:
                self._frames.clear()

# ══════════════════════════════════════════════════════════════════════════════
# QUESTION BANK — aligned to CS/AI/Python tech stack (synopsis Ch.5)
# Categories: hr, technical, behavioural  |  Difficulty: easy, medium, hard
# ══════════════════════════════════════════════════════════════════════════════
QUESTION_BANK = [
    # ── HR / Background ───────────────────────────────────────────────────────
    {"q": "Walk me through your background and the experience most relevant to this role.",
     "category": "hr", "difficulty": "easy"},
    {"q": "Where do you see your career heading in the next three to five years?",
     "category": "hr", "difficulty": "easy"},
    {"q": "Describe a time you explained a complex technical idea to a non-technical person.",
     "category": "hr", "difficulty": "easy"},
    {"q": "What motivates you to keep learning new technologies in AI and software engineering?",
     "category": "hr", "difficulty": "easy"},
    {"q": "How do you stay up to date with developments in AI and machine learning?",
     "category": "hr", "difficulty": "easy"},
    {"q": "Describe the work environment where you do your best engineering work.",
     "category": "hr", "difficulty": "easy"},
    {"q": "What is the most important lesson your projects have taught you so far?",
     "category": "hr", "difficulty": "easy"},

    # ── Behavioural ───────────────────────────────────────────────────────────
    {"q": "Tell me about a project that did not go as planned. What did you learn?",
     "category": "behavioural", "difficulty": "medium"},
    {"q": "Tell me about a time you had to work under pressure or meet a tight deadline.",
     "category": "behavioural", "difficulty": "medium"},
    {"q": "Give an example of when you took initiative on a project without being asked.",
     "category": "behavioural", "difficulty": "medium"},
    {"q": "Tell me about a time you received critical feedback on your code or design. How did you respond?",
     "category": "behavioural", "difficulty": "medium"},
    {"q": "Describe a situation where you had to quickly learn an unfamiliar library or framework to deliver.",
     "category": "behavioural", "difficulty": "medium"},
    {"q": "Give an example of when you had to debug a problem that took much longer than expected.",
     "category": "behavioural", "difficulty": "medium"},
    {"q": "Describe a time you had to choose between two technical approaches. How did you decide?",
     "category": "behavioural", "difficulty": "medium"},
    {"q": "Tell me about a project where you had to integrate multiple external APIs or services.",
     "category": "behavioural", "difficulty": "medium"},
    {"q": "Describe a situation where you had to balance adding features against keeping code maintainable.",
     "category": "behavioural", "difficulty": "hard"},
    {"q": "Tell me about a time you had to optimise a slow piece of code or pipeline. What was your approach?",
     "category": "behavioural", "difficulty": "hard"},

    # ── Python ────────────────────────────────────────────────────────────────
    {"q": "What is the difference between a list, a tuple, and a set in Python? When would you use each?",
     "category": "technical", "difficulty": "easy"},
    {"q": "What are Python decorators and how have you used them in your projects?",
     "category": "technical", "difficulty": "easy"},
    {"q": "What is a Python generator and when is it preferable to a list?",
     "category": "technical", "difficulty": "easy"},
    {"q": "Explain the difference between `*args` and `**kwargs` in Python functions.",
     "category": "technical", "difficulty": "easy"},
    {"q": "What are context managers in Python and how do you implement one with `with`?",
     "category": "technical", "difficulty": "easy"},
    {"q": "What is version control and how do you use Git branches in a project workflow?",
     "category": "technical", "difficulty": "easy"},
    {"q": "Explain how Python's asyncio works and when you would use async/await.",
     "category": "technical", "difficulty": "medium"},
    {"q": "What is the difference between `multiprocessing` and `threading` in Python? When does each help?",
     "category": "technical", "difficulty": "medium"},
    {"q": "How does Python manage memory and what tools can you use to detect memory leaks?",
     "category": "technical", "difficulty": "medium"},

    # ── AI / ML / NLP ─────────────────────────────────────────────────────────
    {"q": "What is the difference between supervised, unsupervised, and reinforcement learning?",
     "category": "technical", "difficulty": "easy"},
    {"q": "What is prompt engineering and what techniques have you used to improve LLM outputs?",
     "category": "technical", "difficulty": "easy"},
    {"q": "Explain what tokenisation is in NLP and why it matters for language models.",
     "category": "technical", "difficulty": "easy"},
    {"q": "What is RAG (Retrieval-Augmented Generation) and what problem does it solve?",
     "category": "technical", "difficulty": "medium"},
    {"q": "Explain what vector embeddings are and how they enable semantic search.",
     "category": "technical", "difficulty": "medium"},
    {"q": "What is a vector database and how does it differ from a traditional relational database?",
     "category": "technical", "difficulty": "medium"},
    {"q": "What is the difference between zero-shot, few-shot, and fine-tuning when working with LLMs?",
     "category": "technical", "difficulty": "medium"},
    {"q": "How would you evaluate the quality and reliability of an LLM-based application?",
     "category": "technical", "difficulty": "medium"},
    {"q": "What are hallucinations in LLMs, why do they occur, and what strategies reduce them?",
     "category": "technical", "difficulty": "medium"},
    {"q": "Explain the transformer architecture — what problem does the attention mechanism solve?",
     "category": "technical", "difficulty": "hard"},
    {"q": "How would you handle rate limits, retries, and fallbacks when calling an LLM API in production?",
     "category": "technical", "difficulty": "hard"},
    {"q": "How would you chunk and index a large document corpus for retrieval in a RAG pipeline?",
     "category": "technical", "difficulty": "hard"},

    # ── APIs, Web & Backend ───────────────────────────────────────────────────
    {"q": "What is a REST API and what are the key principles of RESTful design?",
     "category": "technical", "difficulty": "easy"},
    {"q": "What is FastAPI and why would you choose it over Flask for building an ML backend?",
     "category": "technical", "difficulty": "easy"},
    {"q": "What is Streamlit and what are its strengths and limitations for ML app development?",
     "category": "technical", "difficulty": "easy"},
    {"q": "What is Pydantic and how does it help with data validation in Python APIs?",
     "category": "technical", "difficulty": "medium"},
    {"q": "How do you handle authentication and authorisation in a REST API?",
     "category": "technical", "difficulty": "medium"},
    {"q": "Explain the difference between synchronous and asynchronous request handling in FastAPI.",
     "category": "technical", "difficulty": "medium"},

    # ── Databases & Storage ───────────────────────────────────────────────────
    {"q": "What is the difference between SQL and NoSQL databases? When would you use each?",
     "category": "technical", "difficulty": "easy"},
    {"q": "What is SQLite and what are its advantages and limitations compared to PostgreSQL?",
     "category": "technical", "difficulty": "easy"},
    {"q": "Explain how indexing works in a database and when you would add one.",
     "category": "technical", "difficulty": "medium"},
    {"q": "What is Qdrant and how does it perform similarity search over vector embeddings?",
     "category": "technical", "difficulty": "medium"},
    {"q": "Explain ACID properties in databases and why they matter for application reliability.",
     "category": "technical", "difficulty": "medium"},
    {"q": "How would you decide between storing data in a relational database versus a vector store?",
     "category": "technical", "difficulty": "hard"},

    # ── Cloud, DevOps & Deployment ────────────────────────────────────────────
    {"q": "What is Docker and why is containerisation useful for deploying AI/ML applications?",
     "category": "technical", "difficulty": "easy"},
    {"q": "What is Azure OpenAI Service and how does it differ from the public OpenAI API?",
     "category": "technical", "difficulty": "easy"},
    {"q": "What are environment variables and why should secrets never be hardcoded in code?",
     "category": "technical", "difficulty": "easy"},
    {"q": "What is CI/CD and how would you set up a basic pipeline for a Python project?",
     "category": "technical", "difficulty": "medium"},
    {"q": "How would you monitor errors, latency, and usage in a deployed Streamlit application?",
     "category": "technical", "difficulty": "medium"},
    {"q": "How would you optimise cost and latency when making repeated calls to Azure OpenAI?",
     "category": "technical", "difficulty": "hard"},
]
CATEGORY_LABELS = {"hr": "HR", "technical": "Technical", "behavioural": "Behavioural"}
DIFFICULTY_COLORS = {"easy": "#10B981", "medium": "#F59E0B", "hard": "#EF4444"}
INTERVIEW_QUESTION_COUNT = 15

# Filler words for fluency analysis (synopsis §5.1 — prosody/fluency signals)
FILLER_WORDS = {"um", "uh", "like", "you know", "basically", "literally",
                "actually", "sort of", "kind of", "i mean", "right", "okay so"}

SKILL_KEYWORDS = [
    "python", "java", "javascript", "typescript", "c++", "c#", "sql", "html", "css",
    "react", "node", "express", "django", "flask", "fastapi", "streamlit", "spring",
    "pandas", "numpy", "scikit-learn", "sklearn", "tensorflow", "keras", "pytorch",
    "opencv", "nlp", "machine learning", "deep learning", "computer vision", "llm",
    "rag", "langchain", "openai", "azure", "aws", "gcp", "docker", "kubernetes",
    "git", "github", "linux", "mongodb", "mysql", "postgresql", "sqlite", "firebase",
    "power bi", "tableau", "excel", "rest api", "api", "data analysis", "data science",
]

RESUME_SECTION_ALIASES = {
    "summary": {"summary", "profile", "objective", "career objective", "about"},
    "skills": {"skills", "technical skills", "core skills", "technologies", "tools"},
    "education": {"education", "academic background", "academics", "qualification"},
    "experience": {"experience", "work experience", "professional experience", "internship", "internships", "employment"},
    "projects": {"projects", "academic projects", "personal projects", "project work"},
    "certifications": {"certifications", "certificates", "courses", "achievements"},
}

CSS = """
<style>
@import url('https://fonts.googleapis.com/css2?family=DM+Sans:ital,wght@0,300;0,400;0,500;0,600;0,700;1,400&family=Outfit:wght@400;500;700;800&family=Fira+Code:wght@400;500&display=swap');
*,*::before,*::after{box-sizing:border-box;margin:0;padding:0;}
:root{
  --bg0:#F4F7FB;--bg1:#FFFFFF;--bg2:#F8FAFC;
  --b1:#E2E8F0;--b2:#CBD5E1;
  --t1:#0F172A;--t2:#334155;--t3:#64748B;
  --vl:#6366F1;--vd:#4338CA;--vg:#EEF2FF;--vb:#C7D2FE;
  --green:#10B981;--gg:#D1FAE5;--gb:#6EE7B7;
  --red:#EF4444;--rg:#FEE2E2;--rb:#FCA5A5;
  --amber:#F59E0B;
  --r:16px;--r2:24px;
  --sans:'DM Sans',sans-serif;--heading:'Outfit',sans-serif;--mono:'Fira Code',monospace;
  --shadow:0 20px 40px -10px rgba(0,0,0,.08),0 10px 15px -5px rgba(0,0,0,.03);
}
html,body{background:var(--bg0)!important;font-family:var(--sans);color:var(--t1);overflow-x:hidden;}
.stApp{background:transparent!important;}
#MainMenu,footer,header{visibility:hidden!important;display:none!important;}
.block-container{background:var(--bg1)!important;border-radius:32px;padding:2.5rem 3.5rem!important;margin-top:2rem!important;margin-bottom:2rem!important;max-width:1200px!important;box-shadow:var(--shadow);}
::-webkit-scrollbar{width:4px;}::-webkit-scrollbar-thumb{background:var(--b2);border-radius:4px;}
/* Nav */
.top-nav{height:64px;display:flex;align-items:center;justify-content:space-between;margin-bottom:2.5rem;border-bottom:1px solid var(--b1);}
.nav-logo{display:flex;align-items:center;gap:10px;}
.nav-gem{width:34px;height:34px;border-radius:10px;background:linear-gradient(135deg,var(--vl),#A855F7);display:flex;align-items:center;justify-content:center;font-size:.68rem;font-weight:800;color:#fff;font-family:var(--heading);box-shadow:0 4px 10px rgba(99,102,241,.3);}
.nav-name{font-weight:800;font-size:1.1rem;color:var(--t1);letter-spacing:-.5px;font-family:var(--heading);}
.nav-right{display:flex;align-items:center;gap:.6rem;}
.user-chip{font-size:.82rem;color:var(--t2);font-weight:600;background:var(--bg2);border:1px solid var(--b1);padding:5px 12px;border-radius:99px;}
/* Pills */
.pill{display:inline-flex;align-items:center;gap:4px;padding:5px 12px;border-radius:99px;font-size:.73rem;font-weight:700;white-space:nowrap;}
.p-v{background:var(--vg);color:var(--vd);border:1px solid var(--vb);}
.p-g{background:var(--gg);color:#047857;border:1px solid var(--gb);}
.p-a{background:#FEF3C7;color:#B45309;border:1px solid #FDE68A;}
.p-dim{background:var(--bg2);color:var(--t3);border:1px solid var(--b1);}
.p-r{background:var(--rg);color:#B91C1C;border:1px solid var(--rb);}
/* Buttons */
.stButton>button{background:var(--t1)!important;color:#fff!important;border:none!important;border-radius:12px!important;padding:.72rem 1.5rem!important;font-family:var(--sans)!important;font-weight:700!important;font-size:.92rem!important;width:100%!important;transition:all .25s cubic-bezier(.4,0,.2,1)!important;letter-spacing:.2px!important;}
.stButton>button:hover{background:var(--vl)!important;transform:translateY(-2px)!important;box-shadow:0 8px 20px rgba(99,102,241,.25)!important;}
.stButton>button:disabled{background:var(--b2)!important;transform:none!important;box-shadow:none!important;cursor:not-allowed!important;}
/* Inputs */
div[data-testid="stTextInput"] input{background:var(--bg2)!important;border:1.5px solid var(--b2)!important;border-radius:12px!important;color:var(--t1)!important;font-family:var(--sans)!important;font-size:.95rem!important;padding:.7rem 1rem!important;transition:all .2s;}
div[data-testid="stTextInput"] input:focus{border-color:var(--vl)!important;outline:none!important;box-shadow:0 0 0 4px var(--vg)!important;background:var(--bg1)!important;}
/* File uploader */
div[data-testid="stFileUploader"]{background:var(--bg2)!important;border:2px dashed var(--b2)!important;border-radius:var(--r2)!important;padding:1rem!important;}
div[data-testid="stFileUploader"]:hover{border-color:var(--vl)!important;background:var(--vg)!important;}
/* Audio */
div[data-testid="stAudio"]{background:transparent!important;border-radius:var(--r)!important;margin-bottom:.4rem!important;}
div[data-testid="stAudioInput"]{background:var(--bg2)!important;border:2px dashed var(--b2)!important;border-radius:var(--r2)!important;padding:1rem!important;margin-bottom:.8rem!important;}
/* Cards */
.card{background:var(--bg1);border:1px solid var(--b1);border-radius:var(--r2);padding:1.5rem;margin-bottom:1rem;box-shadow:0 2px 8px rgba(0,0,0,.04);}
.kv-row{display:flex;align-items:center;justify-content:space-between;padding:.45rem 0;border-bottom:1px solid var(--b1);font-size:.84rem;}
.kv-row:last-child{border-bottom:none;}
.kv-key{color:var(--t2);font-weight:600;}
/* Expander */
.stExpander{border:1px solid var(--b1)!important;border-radius:var(--r)!important;background:var(--bg1)!important;overflow:hidden;margin-bottom:1rem!important;}
.stExpander summary{font-weight:600!important;color:var(--t1)!important;background:var(--bg2)!important;padding:.8rem 1rem!important;}
/* Landing */
@keyframes slideUpFade{from{opacity:0;transform:translateY(18px);}to{opacity:1;transform:translateY(0);}}
.landing-hero{padding:1.5rem 3rem 1.5rem 0;animation:slideUpFade .55s ease-out forwards;}
.landing-title{font-size:3.2rem;font-weight:800;font-family:var(--heading);line-height:1.1;margin-bottom:1.2rem;letter-spacing:-1.5px;color:var(--t1);}
.landing-title span{background:linear-gradient(135deg,var(--vl),#A855F7);-webkit-background-clip:text;-webkit-text-fill-color:transparent;}
.feat-badge{display:inline-flex;align-items:center;gap:7px;background:var(--bg1);border:1px solid var(--b1);padding:7px 14px;border-radius:99px;font-size:.82rem;font-weight:600;color:var(--t2);box-shadow:0 2px 6px rgba(0,0,0,.05);margin-right:8px;margin-bottom:8px;}
/* Setup */
.setup-hero{text-align:center;padding:1.5rem 0 2rem;}
.setup-title{font-size:2rem;font-weight:800;font-family:var(--heading);color:var(--t1);letter-spacing:-.5px;margin-bottom:.5rem;}
.setup-sub{font-size:1rem;color:var(--t3);}
.step-badge{display:inline-flex;align-items:center;gap:6px;background:var(--vg);border:1px solid var(--vb);color:var(--vd);padding:4px 12px;border-radius:99px;font-size:.75rem;font-weight:700;margin-bottom:1rem;}
/* Interview */
.q-header{display:flex;align-items:center;justify-content:space-between;margin-bottom:1.4rem;padding-bottom:1rem;border-bottom:1px solid var(--b1);}
.q-counter{font-size:.72rem;font-weight:700;color:var(--t3);letter-spacing:1.2px;text-transform:uppercase;}
.q-box{background:var(--bg1);border:1px solid var(--b1);border-left:4px solid var(--vl);border-radius:var(--r);padding:1.8rem;margin-bottom:.6rem;box-shadow:0 6px 16px rgba(0,0,0,.05);}
.q-text{font-size:1.35rem;font-weight:600;font-family:var(--heading);color:var(--t1);line-height:1.4;letter-spacing:-.3px;}
.q-meta{display:flex;gap:8px;margin-top:.8rem;}
.panel-label{font-size:.68rem;font-weight:700;letter-spacing:1.5px;text-transform:uppercase;color:var(--t3);margin-bottom:.7rem;}
/* Progress bar */
.prog-wrap{background:var(--b1);border-radius:99px;height:6px;margin:1rem 0;overflow:hidden;}
.prog-fill{height:100%;border-radius:99px;background:linear-gradient(90deg,var(--vl),#A855F7);transition:width .5s ease;}
/* Elo badge */
.elo-badge{display:inline-flex;align-items:center;gap:5px;font-size:.72rem;font-weight:700;padding:3px 10px;border-radius:99px;background:#F0E6FF;color:#7C3AED;border:1px solid #DDD6FE;}
/* Evaluating */
.eval-screen{text-align:center;padding:4rem 2rem;}
.eval-icon{font-size:4rem;margin-bottom:1.5rem;animation:pulse 1.5s ease-in-out infinite;}
@keyframes pulse{0%,100%{transform:scale(1);}50%{transform:scale(1.1);}}
/* Results */
@keyframes pulseGradient{0%{background-position:0% 50%;}50%{background-position:100% 50%;}100%{background-position:0% 50%;}}
.results-card{background:linear-gradient(-45deg,#6366F1,#8B5CF6,#10B981,#F59E0B);background-size:300% 300%;animation:pulseGradient 8s ease infinite;border-radius:28px;padding:4px;margin:1.5rem 0;box-shadow:0 20px 40px rgba(99,102,241,.2);}
.results-inner{background:var(--bg1);border-radius:24px;padding:3rem 2rem;text-align:center;}
.res-title{font-size:2.2rem;font-weight:800;font-family:var(--heading);background:linear-gradient(135deg,var(--vl),#A855F7);-webkit-background-clip:text;-webkit-text-fill-color:transparent;margin-bottom:.8rem;letter-spacing:-1px;}
.res-name{font-size:1.1rem;color:var(--t2);font-weight:500;margin-bottom:2rem;}
.res-score{font-size:4.5rem;font-weight:800;line-height:1;font-family:var(--heading);letter-spacing:-2px;}
.res-total{font-size:1.3rem;color:var(--t3);font-weight:600;}
.score-grid{display:flex;gap:1.5rem;justify-content:center;flex-wrap:wrap;margin-top:1.8rem;}
.score-cell{text-align:center;}
.score-val{font-size:1.4rem;font-weight:800;}
.score-lbl{font-size:.68rem;color:var(--t3);font-weight:700;letter-spacing:.5px;text-transform:uppercase;margin-top:2px;}
.eval-feedback{font-size:.88rem;color:var(--t2);line-height:1.6;border-left:3px solid var(--vb);padding-left:12px;margin-top:.8rem;}
.report-box{background:var(--bg2);border:1px solid var(--b1);border-radius:var(--r2);padding:1.8rem;margin-bottom:1rem;}
/* Fluency chip */
.fluency-row{display:flex;gap:8px;flex-wrap:wrap;margin-top:.5rem;}
.fluency-chip{font-size:.72rem;font-weight:600;padding:3px 10px;border-radius:99px;background:var(--bg2);border:1px solid var(--b1);color:var(--t3);}
/* Professional recruiter interface override */
:root{
  --bg0:#F6F7F9;--bg1:#FFFFFF;--bg2:#F9FAFB;
  --b1:#E5E7EB;--b2:#D1D5DB;
  --t1:#111827;--t2:#374151;--t3:#6B7280;
  --vl:#0F766E;--vd:#115E59;--vg:#ECFDF5;--vb:#99F6E4;
  --green:#059669;--gg:#ECFDF5;--gb:#A7F3D0;
  --red:#DC2626;--rg:#FEF2F2;--rb:#FECACA;
  --amber:#D97706;
  --r:8px;--r2:8px;
  --shadow:0 14px 35px rgba(17,24,39,.07);
}
html,body{background:linear-gradient(180deg,#F8FAFC 0%,#F3F4F6 100%)!important;}
.block-container{border-radius:8px!important;border:1px solid var(--b1);box-shadow:var(--shadow);padding:2.25rem 2.75rem!important;margin-top:1.5rem!important;}
.top-nav{height:56px;margin-bottom:2rem;}
.nav-gem{width:32px;height:32px;border-radius:8px;background:#111827;box-shadow:none;letter-spacing:.5px;}
.nav-name{font-size:1rem;letter-spacing:0;}
.user-chip,.pill{border-radius:8px;}
.p-v{background:#F0FDFA;color:#0F766E;border-color:#99F6E4;}
.p-g{background:#ECFDF5;color:#047857;border-color:#A7F3D0;}
.stButton>button{border-radius:8px!important;background:#111827!important;box-shadow:none!important;letter-spacing:0!important;}
.stButton>button:hover{background:#0F766E!important;transform:none!important;box-shadow:0 8px 20px rgba(15,118,110,.16)!important;}
div[data-testid="stTextInput"] input,div[data-testid="stFileUploader"],div[data-testid="stAudioInput"]{border-radius:8px!important;}
div[data-testid="stTextInput"] input:focus{border-color:#0F766E!important;box-shadow:0 0 0 3px rgba(15,118,110,.12)!important;}
div[data-testid="stFileUploader"]:hover{border-color:#0F766E!important;background:#F0FDFA!important;}
.card{border-radius:8px;padding:1.25rem;box-shadow:none;}
.landing-hero{padding:1rem 2rem 1rem 0;}
.landing-title{font-size:2.8rem;letter-spacing:0;line-height:1.08;}
.landing-title span{background:none;-webkit-background-clip:initial;-webkit-text-fill-color:currentColor;color:#0F766E;}
.setup-title{font-size:1.75rem;letter-spacing:0;}
.q-box{border-left:3px solid #0F766E;border-radius:8px;box-shadow:none;background:#FFFFFF;}
.q-text{font-size:1.22rem;letter-spacing:0;}
.prog-fill{background:#0F766E;}
.eval-icon{display:none;}
.results-card{background:#111827;border-radius:8px;animation:none;box-shadow:none;}
.results-inner{border-radius:8px;}
.res-title{background:none;-webkit-background-clip:initial;-webkit-text-fill-color:currentColor;color:#111827;letter-spacing:0;}
.report-box{border-radius:8px;}
</style>
"""

# ══════════════════════════════════════════════════════════════════════════════
# DATABASE  (SQLite — synopsis §4.2 — persistent storage across sessions)
# ══════════════════════════════════════════════════════════════════════════════
DB_PATH = Path(os.getenv("NGR_DB_PATH", "ngr_candidates.db"))

def db_init():
    """Create tables if they don't exist."""
    with sqlite3.connect(DB_PATH) as con:
        con.execute("""
            CREATE TABLE IF NOT EXISTS candidates (
                id        INTEGER PRIMARY KEY AUTOINCREMENT,
                name      TEXT,
                email     TEXT,
                role      TEXT,
                date      TEXT,
                score     REAL,
                answered  INTEGER,
                total     INTEGER,
                elo_final REAL,
                report    TEXT,
                answers   TEXT,
                UNIQUE(email, date)
            )""")
        # Track which question hashes each candidate has seen (synopsis §5.1 — no repeats)
        con.execute("""
            CREATE TABLE IF NOT EXISTS seen_questions (
                email     TEXT,
                q_hash    TEXT,
                PRIMARY KEY (email, q_hash)
            )""")
        con.commit()

def db_get_seen_hashes(email: str) -> set:
    """Return set of question hashes already seen by this candidate."""
    try:
        db_init()
        with sqlite3.connect(DB_PATH) as con:
            rows = con.execute(
                "SELECT q_hash FROM seen_questions WHERE email=?", (email,)
            ).fetchall()
        return {r[0] for r in rows}
    except Exception as e:
        log.warning(f"DB seen_questions load error: {e}")
        return set()

def db_mark_seen(email: str, q_hashes: List[str]):
    """Record that this candidate has seen these question hashes."""
    try:
        db_init()
        with sqlite3.connect(DB_PATH) as con:
            con.executemany(
                "INSERT OR IGNORE INTO seen_questions (email, q_hash) VALUES (?,?)",
                [(email, h) for h in q_hashes]
            )
            con.commit()
    except Exception as e:
        log.warning(f"DB mark_seen error: {e}")

def db_save(record: Dict):
    """Insert or replace a candidate record."""
    try:
        db_init()
        with sqlite3.connect(DB_PATH) as con:
            con.execute("""
                INSERT OR REPLACE INTO candidates
                (name, email, role, date, score, answered, total, elo_final, report, answers)
                VALUES (?,?,?,?,?,?,?,?,?,?)""",
                (record["name"], record["email"], record["role"], record["date"],
                 record["score"], record["answered"], record["total"],
                 record["elo_final"], record["report"],
                 json.dumps(record["answers"])))
            con.commit()
    except Exception as e:
        log.warning(f"DB save error: {e}")

def db_load_all() -> List[Dict]:
    """Return all candidate records sorted by score desc."""
    try:
        db_init()
        with sqlite3.connect(DB_PATH) as con:
            rows = con.execute(
                "SELECT name,email,role,date,score,answered,total,elo_final,report,answers "
                "FROM candidates ORDER BY score DESC"
            ).fetchall()
        result = []
        for r in rows:
            result.append({
                "name": r[0], "email": r[1], "role": r[2], "date": r[3],
                "score": r[4], "answered": r[5], "total": r[6],
                "elo_final": r[7], "report": r[8],
                "answers": json.loads(r[9] or "[]"),
            })
        return result
    except Exception as e:
        log.warning(f"DB load error: {e}")
        return []

def db_delete(email: str, date: str):
    """Delete a candidate record by email + date."""
    try:
        with sqlite3.connect(DB_PATH) as con:
            con.execute("DELETE FROM candidates WHERE email=? AND date=?", (email, date))
            con.commit()
    except Exception as e:
        log.warning(f"DB delete error: {e}")

# ── Secrets helper ─────────────────────────────────────────────────────────
def _safe_secret(key: str) -> Optional[str]:
    try: return st.secrets[key]
    except Exception: return None

# ── AI Client ──────────────────────────────────────────────────────────────
@st.cache_resource
def get_ai_client():
    if not _OAI: return None, None
    groq_key = os.getenv("GROQ_API_KEY") or _safe_secret("GROQ_API_KEY")
    if groq_key:
        try:
            c = OpenAI(api_key=groq_key, base_url="https://api.groq.com/openai/v1")
            return c, ("groq", "llama-3.3-70b-versatile")
        except Exception as e: log.warning(f"Groq init failed: {e}")
    azure_key      = os.getenv("AZURE_OPENAI_API_KEY")  or _safe_secret("AZURE_OPENAI_API_KEY")
    azure_endpoint = os.getenv("AZURE_OPENAI_ENDPOINT") or _safe_secret("AZURE_OPENAI_ENDPOINT")
    azure_deploy   = os.getenv("AZURE_OPENAI_DEPLOYMENT") or _safe_secret("AZURE_OPENAI_DEPLOYMENT") or "gpt-4o"
    if azure_key and azure_endpoint:
        try:
            c = AzureOpenAI(api_key=azure_key, azure_endpoint=azure_endpoint, api_version="2024-06-01")
            return c, ("azure", azure_deploy)
        except Exception as e: log.warning(f"Azure init failed: {e}")
    oai_key = os.getenv("OPENAI_API_KEY") or _safe_secret("OPENAI_API_KEY")
    if oai_key:
        try: return OpenAI(api_key=oai_key), ("openai", "gpt-4o-mini")
        except Exception as e: log.warning(f"OpenAI init failed: {e}")
    return None, None

def chat_complete(messages: list, max_tokens: int = 800, temperature: float = 0.7) -> str:
    client, info = get_ai_client()
    if not client or not info: return ""
    _, model = info
    try:
        r = client.chat.completions.create(model=model, messages=messages, max_tokens=max_tokens, temperature=temperature)
        return (r.choices[0].message.content or "").strip() if r.choices else ""
    except Exception as e:
        log.warning(f"chat_complete failed: {e}"); return ""

def _find_ffmpeg() -> Optional[str]:
    """Find ffmpeg executable — checks PATH, env var, imageio, and common locations."""
    import importlib, shutil, subprocess
    # 1. shutil.which covers PATH on all platforms
    found = shutil.which("ffmpeg")
    if found: return found
    # 2. imageio-ffmpeg (installed via pip install imageio[ffmpeg])
    try:
        imageio_ffmpeg = importlib.import_module("imageio_ffmpeg")
        exe = imageio_ffmpeg.get_ffmpeg_exe()
        if exe: return exe
    except Exception: pass
    try:
        imageio = importlib.import_module("imageio")
        exe = imageio.plugins.ffmpeg.get_exe()
        if exe: return exe
    except Exception: pass
    # 3. Explicit path candidates
    candidates = [
        os.getenv("FFMPEG_PATH", ""),
        r"C:\Users\Admin\Downloads\ffmpeg-8.0.1-essentials_build\bin\ffmpeg.exe",
        r"C:\Users\DELL-IN\Downloads\ffmpeg-8.0.1-essentials_build\bin\ffmpeg.exe",
        r"C:\Users\Admin\anaconda3\pkgs\imageio-2.26.0-py310haa95532_0\Lib\site-packages\imageio\plugins\ffmpeg.exe",
        r"C:\ffmpeg\bin\ffmpeg.exe",
        r"C:\Program Files\ffmpeg\bin\ffmpeg.exe",
        r"C:\ProgramData\chocolatey\bin\ffmpeg.exe",
        "/usr/bin/ffmpeg",
        "/usr/local/bin/ffmpeg",
    ]
    for c in candidates:
        if not c: continue
        try:
            r = subprocess.run([c, "-version"], capture_output=True, timeout=5)
            if r.returncode == 0: return c
        except Exception: continue
    return None

def _to_mp3(raw: bytes, in_fmt: str = "wav") -> Optional[bytes]:
    import subprocess, tempfile, os as _os
    ffmpeg_exe = _find_ffmpeg()
    if not ffmpeg_exe:
        log.warning("ffmpeg not found — install ffmpeg and add it to PATH, or set FFMPEG_PATH env var")
        return None
    inp = out = ""
    try:
        with tempfile.NamedTemporaryFile(suffix=f".{in_fmt}", delete=False) as f:
            f.write(raw); inp = f.name
        out = inp + ".mp3"
        r = subprocess.run(
            [ffmpeg_exe, "-y", "-i", inp,
             "-vn", "-ar", "16000", "-ac", "1", "-b:a", "64k", out],
            capture_output=True, timeout=30
        )
        if r.returncode != 0:
            log.warning(f"ffmpeg failed: {r.stderr.decode(errors='ignore')[:300]}")
            return None
        if _os.path.exists(out) and _os.path.getsize(out) > 100:
            with open(out, "rb") as f: return f.read()
    except Exception as e:
        log.warning(f"ffmpeg error: {e}")
    finally:
        for p in (inp, out):
            try:
                if p and _os.path.exists(p): _os.unlink(p)
            except: pass
    return None

# ── Transcription ──────────────────────────────────────────────────────────
def do_transcribe(raw: bytes, fmt: str = "wav") -> Tuple[str, float]:
    if not raw or len(raw) < 100: return "", 0.0
    mp3 = _to_mp3(raw, fmt)
    if not mp3:
        log.warning("do_transcribe: mp3 conversion failed")
        return "", 0.0

    # Only reject obvious hallucinations — do NOT reject short answers
    HALLUCINATION_SIGNALS = ["www.", "http", "subtitle", "subscrib", "zeoranger",
                              "opensubtitles", ".co.uk", "thank you for watching"]
    def is_hallucination(text: str) -> bool:
        t = text.lower().strip()
        if not t: return True
        return any(s in t for s in HALLUCINATION_SIGNALS)

    groq_key = os.getenv("GROQ_API_KEY") or _safe_secret("GROQ_API_KEY")
    if groq_key:
        try:
            import requests
            resp = requests.post(
                "https://api.groq.com/openai/v1/audio/transcriptions",
                headers={"Authorization": f"Bearer {groq_key}"},
                files={"file": ("audio.mp3", mp3, "audio/mpeg")},
                data={"model": "whisper-large-v3-turbo", "response_format": "text"},
                timeout=60
            )
            if resp.status_code == 200:
                text = resp.text.strip()
                if text and not is_hallucination(text):
                    return text, 0.0
                log.warning(f"Groq hallucination detected: {repr(text[:80])}")
            else:
                log.warning(f"Groq HTTP {resp.status_code}: {resp.text[:200]}")
        except Exception as e:
            log.warning(f"Groq transcription error: {e}")

    azure_key     = os.getenv("AZURE_OPENAI_API_KEY")  or _safe_secret("AZURE_OPENAI_API_KEY")
    azure_endpoint = os.getenv("AZURE_OPENAI_ENDPOINT") or _safe_secret("AZURE_OPENAI_ENDPOINT")
    whisper_model  = os.getenv("AZURE_OPENAI_WHISPER_DEPLOYMENT") or _safe_secret("AZURE_OPENAI_WHISPER_DEPLOYMENT") or "whisper"
    if azure_key and azure_endpoint:
        try:
            import requests
            base = re.match(r"(https?://[^/]+)", azure_endpoint.strip())
            base_url = base.group(1).rstrip("/") if base else azure_endpoint.rstrip("/")
            url = f"{base_url}/openai/deployments/{whisper_model}/audio/transcriptions?api-version=2024-06-01"
            resp = requests.post(
                url,
                headers={"api-key": azure_key},
                files={"file": ("audio.mp3", mp3, "audio/mpeg")},
                data={"response_format": "text"},
                timeout=60
            )
            if resp.status_code == 200:
                text = resp.text.strip()
                if text and not is_hallucination(text):
                    return text, 0.0
                log.warning(f"Azure hallucination detected: {repr(text[:80])}")
            else:
                log.warning(f"Azure Whisper HTTP {resp.status_code}: {resp.text[:200]}")
        except Exception as e:
            log.warning(f"Azure transcription error: {e}")

    log.warning("do_transcribe: all methods failed")
    return "", 0.0

# ── Prosody analysis via Librosa (synopsis §5.1 — pitch, energy, pauses) ──
def analyse_prosody(raw_audio: bytes, fmt: str = "webm") -> Dict:
    """
    Real prosody analysis using Librosa (synopsis §5.2 — speech analysis).
    Extracts: speech rate proxy, pitch mean/std, RMS energy, pause ratio.
    Returns a prosody dict merged into the fluency result.
    Falls back to empty dict if librosa unavailable.
    """
    try:
        import librosa, tempfile, subprocess, os as _os
        import numpy as np

        # Convert to WAV via ffmpeg for librosa
        ffmpeg_exe = _find_ffmpeg()
        if not ffmpeg_exe:
            return {}

        inp = out = ""
        try:
            with tempfile.NamedTemporaryFile(suffix=f".{fmt}", delete=False) as f:
                f.write(raw_audio); inp = f.name
            out = inp + "_prosody.wav"
            r = subprocess.run(
                [ffmpeg_exe, "-y", "-i", inp, "-ar", "22050", "-ac", "1", out],
                capture_output=True, timeout=30
            )
            if r.returncode != 0 or not _os.path.exists(out):
                return {}

            y, sr = librosa.load(out, sr=22050)
            duration = librosa.get_duration(y=y, sr=sr)
            if duration < 0.5:
                return {}

            # ── Pitch (F0) via piptrack ───────────────────────────────────────
            pitches, magnitudes = librosa.piptrack(y=y, sr=sr)
            pitch_values = pitches[magnitudes > np.median(magnitudes[magnitudes > 0]) * 0.5]
            pitch_values = pitch_values[(pitch_values > 50) & (pitch_values < 500)]
            pitch_mean = round(float(np.mean(pitch_values)), 1) if len(pitch_values) > 0 else 0.0
            pitch_std  = round(float(np.std(pitch_values)),  1) if len(pitch_values) > 0 else 0.0

            # ── Energy (RMS) ──────────────────────────────────────────────────
            rms = librosa.feature.rms(y=y)[0]
            energy_mean = float(round(float(np.mean(rms)), 4))

            # ── Pause detection (low-energy frames = silence) ─────────────────
            silence_threshold = np.max(rms) * 0.05
            silent_frames = np.sum(rms < silence_threshold)
            total_frames  = len(rms)
            pause_ratio   = float(round(silent_frames / max(total_frames, 1), 3))

            # ── Speech rate proxy: words per second ───────────────────────────
            # (word_count filled in by caller after transcript analysis)
            speech_rate_proxy = float(round(1.0 - pause_ratio, 3))  # 0-1, higher = more speech

            # ── Prosody score: penalise high pauses, reward moderate pitch var ─
            pause_penalty  = pause_ratio * 4          # 0→4 pts off
            pitch_var_bonus = min(pitch_std / 50, 1) * 2  # up to 2 pts for expressive pitch
            prosody_score  = max(0, min(10, round(8 - pause_penalty + pitch_var_bonus, 1)))

            return {
                "pitch_mean":      pitch_mean,
                "pitch_std":       pitch_std,
                "energy_mean":     energy_mean,
                "pause_ratio":     pause_ratio,
                "speech_rate":     speech_rate_proxy,
                "prosody_score":   prosody_score,
                "duration_sec":    round(duration, 1),
            }
        finally:
            for p in (inp, out):
                try:
                    if p and _os.path.exists(p): _os.unlink(p)
                except: pass

    except ImportError:
        log.warning("librosa not installed — prosody analysis skipped")
        return {}
    except Exception as e:
        log.warning(f"Prosody analysis error: {e}")
        return {}

def analyse_fluency(transcript: str, raw_audio: bytes = None, fmt: str = "wav") -> dict:
    import re

    if not transcript:
        return {}

    text = transcript.lower().strip()

    # Words
    words = re.findall(r'\b\w+\b', text)
    word_count = len(words)

    # Fillers
    fillers_list = ["um", "uh", "hmm", "like", "you know", "basically", "actually"]
    filler_count = 0
    fillers_found = []

    for f in fillers_list:
        count = len(re.findall(rf'\b{re.escape(f)}\b', text))
        if count:
            filler_count += count
            fillers_found.append(f)

    filler_rate = filler_count / max(word_count, 1)

    # Duration estimate
    duration_sec = max(word_count / 2.5, 1.0)

    # Pause estimation
    pause_marks = len(re.findall(r'[.,;!?]', transcript))
    pause_ratio = pause_marks / max(word_count, 1)

    # Prosody
    prosody_score = max(0, 10 - pause_ratio * 10)

    return {
        "word_count": word_count,
        "filler_count": filler_count,
        "filler_rate": round(filler_rate, 3),
        "fillers_found": fillers_found,
        "duration_sec": round(duration_sec, 2),
        "pause_ratio": round(pause_ratio, 3),
        "pitch_mean": 0.0,
        "pitch_std": 0.0,
        "prosody_score": round(prosody_score, 2)
    }

def enhanced_fluency_score(fluency: dict) -> float:
    words = fluency.get("word_count", 0)
    fillers = fluency.get("filler_count", 0)
    duration = fluency.get("duration_sec", 1.0)
    pause_ratio = fluency.get("pause_ratio", 0.0)

    words_per_sec = words / max(duration, 1)

    # Ideal speaking rate
    rate_score = max(0, 10 - abs(words_per_sec - 2.5) * 3)

    filler_score = max(0, 10 - fillers)
    pause_score = max(0, 10 - pause_ratio * 10)

    final_score = (
        0.4 * rate_score +
        0.3 * filler_score +
        0.3 * pause_score
    )

    return round(final_score, 2)

def q_hash(q: Dict) -> str:
    """Stable hash for a question dict — used to track seen questions across sessions."""
    import hashlib
    return hashlib.md5(q["q"].encode()).hexdigest()[:16]

def pick_next_question(question_pool: List[Dict], elo: float,
                       used_indices: set, seen_hashes: set) -> Optional[Dict]:
    """
    Adaptive question selection (synopsis §5.1 — Elo-based difficulty).
    Priority order:
      1. Match Elo difficulty AND not seen this session AND not seen in past sessions
      2. Relax difficulty — any unseen question (this + past sessions)
      3. Relax past-session constraint — unseen this session only
      4. Last resort — anything not used this session
    """
    if elo < 1200:   target = "easy"
    elif elo < 1500: target = "medium"   # raised from 1400 → 1500
    else:            target = "hard"     # needs ~10 strong answers to reach

    def unseen_pool(relax_past=False):
        result = []
        for i, q in enumerate(question_pool):
            if i in used_indices: continue                           # used this session
            qh = q_hash(q)
            if not relax_past and qh in seen_hashes: continue       # seen in past sessions
            result.append((i, q))
        return result

    pool = unseen_pool(relax_past=False)
    by_diff = [(i, q) for i, q in pool if q["difficulty"] == target]
    if by_diff:
        i, q = random.choice(by_diff)
    elif pool:
        i, q = random.choice(pool)
    else:
        # All questions seen before — fall back, ignore past sessions
        pool2 = unseen_pool(relax_past=True)
        by_diff2 = [(i, q) for i, q in pool2 if q["difficulty"] == target]
        if by_diff2:   i, q = random.choice(by_diff2)
        elif pool2:    i, q = random.choice(pool2)
        else: return None

    return q

# ── Question generation ────────────────────────────────────────────────────
def _context_value(text: str, label: str) -> str:
    match = re.search(rf"^{re.escape(label)}:\s*(.+)$", text, flags=re.I | re.M)
    return match.group(1).strip() if match else ""

def _split_context_items(value: str) -> List[str]:
    items = []
    for part in re.split(r"\s+\|\s+|,\s*", value or ""):
        part = re.sub(r"\s+", " ", part).strip(" .;-")
        if part and part.lower() not in {"not clearly listed", "unknown"}:
            items.append(part[:120])
    return list(dict.fromkeys(items))

def _resume_question_signals(text: str, role: str) -> Dict[str, List[str]]:
    skills = _split_context_items(_context_value(text, "Skills"))
    projects = _split_context_items(_context_value(text, "Projects"))
    experience = _split_context_items(_context_value(text, "Experience"))

    if not skills:
        skills = _extract_skills(text, _split_resume_sections(_clean_resume_lines(text)))
    if not projects:
        lines = _clean_resume_lines(text)
        projects = _top_section_items(
            [line for line in lines if re.search(r"\b(project|built|developed|implemented|created)\b", line, re.I)],
            3,
        )

    role_terms = [part.strip() for part in re.split(r"[/,|-]", role or "") if part.strip()]
    return {
        "skills": skills[:8],
        "projects": projects[:4],
        "experience": experience[:3],
        "role_terms": role_terms[:3],
    }

def _first_or(items: List[str], fallback: str) -> str:
    return items[0] if items else fallback

def _short_label(value: str, fallback: str) -> str:
    value = re.split(r"\s[-–—:]\s", value or "", maxsplit=1)[0].strip()
    return value[:80] if value else fallback

def _dedupe_questions(questions: List[Dict], n: int) -> List[Dict]:
    seen = set()
    result = []
    for q in questions:
        key = re.sub(r"\W+", " ", q["q"].lower()).strip()
        if key in seen:
            continue
        seen.add(key)
        result.append(q)
        if len(result) >= n:
            break
    return result

def do_gen_questions(text: str, role: str, n: int = 5) -> List[Dict]:
    """Generate resume-specific questions locally from parsed resume signals."""
    signals = _resume_question_signals(text, role)
    skills = signals["skills"]
    projects = signals["projects"]
    experience = signals["experience"]
    role_name = role.strip() or "this role"
    primary_skill = _first_or(skills, "your strongest technical skill")
    secondary_skill = skills[1] if len(skills) > 1 else primary_skill
    project = _short_label(_first_or(projects, ""), "one of your main projects")
    experience_item = _short_label(_first_or(experience, ""), "your recent technical experience")

    questions = [
        {
            "q": (
                f"In {project}, walk me through the end-to-end architecture you built for a {role_name} context. "
                f"Which components were hardest to design, and what trade-offs did you make?"
            ),
            "category": "technical",
            "difficulty": "hard",
            "personalised": True,
        },
        {
            "q": (
                f"You listed {primary_skill}. Explain a technically deep problem you solved with it, "
                f"including the data structures, libraries, APIs, or design patterns you chose and why."
            ),
            "category": "technical",
            "difficulty": "hard",
            "personalised": True,
        },
        {
            "q": (
                f"If {project} started failing in production with slow responses or incorrect outputs, "
                f"how would you debug it step by step, measure the root cause, and verify the fix?"
            ),
            "category": "technical",
            "difficulty": "hard",
            "personalised": True,
        },
        {
            "q": (
                f"Compare {primary_skill} and {secondary_skill} in the context of {role_name}. "
                f"Where would each fit in a real system, and what mistakes should an engineer avoid?"
            ),
            "category": "technical",
            "difficulty": "medium",
            "personalised": True,
        },
        {
            "q": (
                f"From {experience_item}, choose one feature or module you worked on. "
                f"How would you improve its reliability, testing strategy, and maintainability today?"
            ),
            "category": "technical",
            "difficulty": "medium",
            "personalised": True,
        },
    ]

    lower_skills = " ".join(skills).lower()
    if any(term in lower_skills for term in ["machine learning", "deep learning", "nlp", "computer vision", "tensorflow", "pytorch", "scikit-learn", "openai", "llm", "rag"]):
        questions.extend([
            {
                "q": (
                    f"For the AI/ML parts of {project}, how would you evaluate model or LLM output quality? "
                    f"Describe the metrics, test set, failure cases, and human review process you would use."
                ),
                "category": "technical",
                "difficulty": "hard",
                "personalised": True,
            },
            {
                "q": (
                    f"If your {primary_skill}-based pipeline produced biased, hallucinated, or unstable results, "
                    f"what safeguards would you add before deploying it to users?"
                ),
                "category": "technical",
                "difficulty": "hard",
                "personalised": True,
            },
        ])

    if any(term in lower_skills for term in ["fastapi", "flask", "django", "streamlit", "react", "node", "rest api", "api"]):
        questions.extend([
            {
                "q": (
                    f"Design the backend/API flow for {project}. What endpoints, validation, error handling, "
                    f"and security checks would you include for a production version?"
                ),
                "category": "technical",
                "difficulty": "hard",
                "personalised": True,
            },
            {
                "q": (
                    f"How would you structure tests for the user-facing and API parts of {project}, "
                    f"including edge cases, failed external services, and bad input?"
                ),
                "category": "technical",
                "difficulty": "medium",
                "personalised": True,
            },
        ])

    if any(term in lower_skills for term in ["sql", "sqlite", "mysql", "postgresql", "mongodb", "firebase"]):
        questions.append({
            "q": (
                f"Explain how you would model and query the data for {project}. "
                f"What indexes, constraints, and data validation would matter as usage grows?"
            ),
            "category": "technical",
            "difficulty": "medium",
            "personalised": True,
        })

    for skill in skills[:6]:
        questions.extend([
            {
                "q": (
                    f"Give a concrete example of how you used {skill}. What was the input, what processing happened, "
                    f"what output was produced, and how did you validate correctness?"
                ),
                "category": "technical",
                "difficulty": "medium",
                "personalised": True,
            },
            {
                "q": (
                    f"What are two non-obvious limitations or failure modes of using {skill} in a real {role_name} project, "
                    f"and how would you reduce those risks?"
                ),
                "category": "technical",
                "difficulty": "hard",
                "personalised": True,
            },
        ])

    for item in projects[:3]:
        project_label = _short_label(item, project)
        questions.extend([
            {
                "q": (
                    f"For {project_label}, describe the most important module in detail. "
                    f"What functions, classes, database tables, or services would exist inside it?"
                ),
                "category": "technical",
                "difficulty": "medium",
                "personalised": True,
            },
            {
                "q": (
                    f"How would you redesign {project_label} if it had to support 10x more users or data? "
                    f"Discuss bottlenecks, caching, database choices, and monitoring."
                ),
                "category": "technical",
                "difficulty": "hard",
                "personalised": True,
            },
        ])

    questions.extend([
        {
            "q": (
                f"Suppose you join as a {role_name} and inherit this codebase. "
                f"What technical checks would you perform in the first day to assess code quality, security, and reliability?"
            ),
            "category": "technical",
            "difficulty": "medium",
            "personalised": True,
        },
        {
            "q": (
                f"How would you design a clean error-handling and logging strategy for {project}, "
                f"especially around file uploads, external APIs, and user input?"
            ),
            "category": "technical",
            "difficulty": "medium",
            "personalised": True,
        },
        {
            "q": (
                f"Explain how you would write automated tests for {project}. "
                f"Which unit tests, integration tests, and edge-case tests would give you the most confidence?"
            ),
            "category": "technical",
            "difficulty": "medium",
            "personalised": True,
        },
    ])

    fallback = [
        dict(q, personalised=True)
        for q in QUESTION_BANK
        if q["category"] == "technical" and q["difficulty"] in {"medium", "hard"}
    ]
    return _dedupe_questions(questions + fallback, n)

# ── Answer evaluation (local NLP scoring: keywords + similarity + structure) ─
STOPWORDS = {
    "a", "an", "and", "are", "as", "at", "be", "by", "for", "from", "how", "i",
    "in", "is", "it", "of", "on", "or", "that", "the", "this", "to", "was",
    "were", "what", "when", "where", "which", "who", "why", "with", "would",
    "you", "your", "we", "our", "they", "their", "them", "then", "than",
}
TECHNICAL_TERMS = {
    "api", "architecture", "async", "authentication", "authorization", "cache",
    "class", "cloud", "complexity", "database", "debug", "deployment", "docker",
    "endpoint", "error", "exception", "index", "latency", "logging", "model",
    "module", "monitoring", "pipeline", "query", "queue", "schema", "security",
    "server", "service", "testing", "validation", "vector", "workflow", "python",
    "sql", "streamlit", "fastapi", "react", "node", "aws", "azure", "docker",
    "kubernetes", "pandas", "numpy", "etl", "nlp", "bert", "vader", "ocr",
    "upload", "pdf", "docx", "mime", "file", "parser", "unit", "integration",
    "external", "fallback", "protected", "scanned", "text", "extraction",
}
POSITIVE_WORDS = {"improved", "optimized", "reliable", "clear", "successful", "efficient", "accurate", "robust"}
NEGATIVE_WORDS = {"failed", "confusing", "wrong", "bad", "broken", "unclear", "unstable", "slow"}
HEDGE_WORDS = {"maybe", "probably", "guess", "basically", "kind", "sort", "might", "maybe"}
KEYWORD_CANONICAL = {
    "api": "api", "backend": "api", "endpoint": "api", "endpoints": "api", "service": "api",
    "validate": "validation", "validates": "validation", "validated": "validation", "validating": "validation",
    "handle": "handling", "handles": "handling", "handled": "handling", "handling": "handling",
    "error": "error", "errors": "error", "exception": "error", "exceptions": "error", "fallback": "error",
    "secure": "security", "security": "security", "unsafe": "security", "auth": "security",
    "test": "testing", "tests": "testing", "tested": "testing", "testing": "testing",
    "parse": "parsing", "parser": "parsing", "parsers": "parsing", "parsed": "parsing", "parsing": "parsing",
    "log": "logging", "logs": "logging", "logged": "logging", "logging": "logging",
    "monitor": "monitoring", "monitors": "monitoring", "monitoring": "monitoring",
    "scale": "scaling", "scaled": "scaling", "scalable": "scaling", "scaling": "scaling",
}

def _nlp_tokens(text: str) -> List[str]:
    return [
        t.lower()
        for t in re.findall(r"[a-zA-Z][a-zA-Z0-9+#.-]*", text or "")
        if len(t) > 1
    ]

def _keywords(text: str) -> List[str]:
    tokens = _nlp_tokens(text)
    result = []
    for token in tokens:
        token = KEYWORD_CANONICAL.get(token, token)
        if token.endswith("ing") and len(token) > 5 and token not in {"handling", "testing", "parsing", "logging", "monitoring", "scaling"}:
            token = token[:-3]
        elif token.endswith("ed") and len(token) > 4:
            token = token[:-2]
        elif token.endswith("s") and len(token) > 4 and token not in {"class", "analysis"}:
            token = token[:-1]
        token = KEYWORD_CANONICAL.get(token, token)
        if token in STOPWORDS:
            continue
        if len(token) <= 2 and token not in {"ai", "ml", "db"}:
            continue
        result.append(token)
    return list(dict.fromkeys(result))

def _keyword_metrics(question: str, answer: str) -> Dict:
    q_keys = set(_keywords(question))
    a_keys = set(_keywords(answer))
    overlap = sorted(q_keys & a_keys)
    coverage = len(overlap) / max(len(q_keys), 1)
    density = len(overlap) / max(len(a_keys), 1)
    score = min(10.0, round(coverage * 12.0 + density * 2.0 + min(len(overlap), 8) * 0.35, 1))
    return {
        "question_keywords": sorted(q_keys)[:18],
        "matched_keywords": overlap[:18],
        "keyword_coverage": round(coverage, 3),
        "keyword_score": score,
    }

def _semantic_similarity(question: str, answer: str) -> float:
    q = set(_keywords(question))
    a = set(_keywords(answer))
    lexical = len(q & a) / max(len(q | a), 1)
    coverage = len(q & a) / max(len(q), 1)
    try:
        from sklearn.feature_extraction.text import TfidfVectorizer
        from sklearn.metrics.pairwise import cosine_similarity
        vectorizer = TfidfVectorizer(stop_words="english", ngram_range=(1, 2), max_features=500)
        matrix = vectorizer.fit_transform([question, answer])
        tfidf = float(cosine_similarity(matrix[0:1], matrix[1:2])[0][0])
        return round(max(tfidf, lexical, coverage * 0.65), 3)
    except Exception:
        return round(max(lexical, coverage * 0.65), 3)

def _structure_score(answer: str) -> Dict:
    words = _nlp_tokens(answer)
    sentence_count = len([s for s in re.split(r"[.!?]+", answer) if len(s.split()) >= 3])
    connector_count = len(re.findall(
        r"\b(first|second|third|then|next|because|therefore|however|for example|for instance|finally|also|while|whereas|so)\b",
        answer, flags=re.I,
    ))
    example_count = len(re.findall(r"\b(example|instance|used|built|implemented|designed|created|developed|deployed)\b", answer, flags=re.I))
    specificity_count = len(re.findall(r"\b\d+%?|\b(api|database|table|endpoint|model|class|function|test|index|cache|queue|service)\b", answer, flags=re.I))

    length_score = min(3.0, len(words) / 25)
    sentence_score = min(2.0, sentence_count * 0.8)
    flow_score = min(2.0, connector_count * 0.45)
    example_score = min(1.5, example_count * 0.35)
    specificity_score = min(1.5, specificity_count * 0.25)
    score = round(min(10.0, length_score + sentence_score + flow_score + example_score + specificity_score), 1)
    return {
        "structure_score": score,
        "sentence_count": sentence_count,
        "connector_count": connector_count,
        "example_count": example_count,
        "specificity_count": specificity_count,
    }

def _technical_depth_score(answer: str, keyword_score: float, semantic_score: float) -> Dict:
    tokens = set(_nlp_tokens(answer))
    tech_hits = sorted(tokens & TECHNICAL_TERMS)
    action_hits = re.findall(
        r"\b(design|debug|measure|validate|optimize|scale|deploy|monitor|test|index|cache|secure|refactor|integrate)\w*\b",
        answer, flags=re.I,
    )
    tradeoff_hits = re.findall(r"\b(trade[- ]?off|bottleneck|constraint|latency|failure|edge case|risk|fallback)\b", answer, flags=re.I)
    score = (
        min(4.0, len(tech_hits) * 0.55)
        + min(2.5, len(action_hits) * 0.45)
        + min(1.5, len(tradeoff_hits) * 0.5)
        + min(1.3, keyword_score * 0.15)
        + min(0.7, semantic_score * 2.0)
    )
    return {"technical_depth": round(min(10.0, score), 1), "technical_terms": tech_hits[:12]}

def _sentiment(answer: str) -> str:
    tokens = set(_nlp_tokens(answer))
    pos = len(tokens & POSITIVE_WORDS)
    neg = len(tokens & NEGATIVE_WORDS)
    if pos > neg:
        return "positive"
    if neg > pos + 1:
        return "negative"
    return "neutral"

def compute_confidence(fluency: dict, eval_data: dict) -> float:
    base = eval_data.get("confidence", 5)

    pause_penalty = fluency.get("pause_ratio", 0) * 5
    filler_penalty = fluency.get("filler_count", 0) * 0.3

    score = base - pause_penalty - filler_penalty
    return max(0, min(10, score))

def _confidence_score(answer: str, fluency: Dict, structure: Dict) -> float:
    words = _nlp_tokens(answer)
    hedge_count = sum(1 for t in words if t in HEDGE_WORDS)
    filler_rate = float(fluency.get("filler_rate", 0.0) or 0.0)
    fluency_score = float(fluency.get("fluency_score", 5) or 5)
    base = 4.0 + min(2.0, len(words) / 55) + min(2.0, structure["specificity_count"] * 0.25)
    base += min(2.0, fluency_score * 0.2)
    base -= min(2.5, hedge_count * 0.45 + filler_rate * 10)
    return round(max(0.0, min(10.0, base)), 1)
def compute_confidence(fluency: dict, eval_data: dict) -> float:
    base = eval_data.get("confidence", 5)

    pause_penalty = fluency.get("pause_ratio", 0) * 5
    filler_penalty = fluency.get("filler_count", 0) * 0.3

    score = base - pause_penalty - filler_penalty
    return max(0, min(10, round(score, 2)))

def do_evaluate_answer(question, transcript, fluency):
    if not transcript or transcript in ["No answer provided", "Skipped"]:
        return {
            "technical_depth": 0,
            "clarity": 0,
            "confidence": 0,
            "relevance": 0,
            "behavioral": 0,
            "sentiment": "neutral",
            "feedback": "No answer provided."
        }

    text = transcript.lower()

    # ── TECHNICAL DEPTH ─────────────────
    tech_keywords = ["algorithm", "data", "complexity", "api", "database", "model", "system"]
    tech_score = min(10, sum(1 for w in tech_keywords if w in text) * 1.5)

    # ── CLARITY ────────────────────────
    word_count = len(text.split())
    clarity = min(10, word_count / 15)

    # ── CONFIDENCE ─────────────────────
    filler_rate = fluency.get("filler_rate", 0)
    confidence = max(0, 10 - filler_rate * 20)

    # ── RELEVANCE ──────────────────────
    relevance = 8 if len(text) > 20 else 4

    # ── 🔥 BEHAVIORAL SCORE (NEW) ──────
    behavioral_keywords = [
        "team", "challenge", "conflict", "lead", "responsibility",
        "managed", "collaborated", "improved", "decision",
        "problem", "solution", "result", "experience"
    ]

    behavioral_hits = sum(1 for w in behavioral_keywords if w in text)

    # STAR method detection (Situation, Task, Action, Result)
    star_score = 0
    if "situation" in text or "when" in text:
        star_score += 2
    if "task" in text or "responsibility" in text:
        star_score += 2
    if "action" in text or "did" in text:
        star_score += 3
    if "result" in text or "outcome" in text:
        star_score += 3

    behavioral = min(10, behavioral_hits + star_score)

    # ── SENTIMENT ──────────────────────
    sentiment = "positive" if confidence > 6 else "neutral"

    return {
        "technical_depth": round(tech_score, 1),
        "clarity": round(clarity, 1),
        "confidence": round(confidence, 1),
        "relevance": round(relevance, 1),
        "behavioral": round(behavioral, 1),   # ✅ IMPORTANT
        "sentiment": sentiment,
        "feedback": "Good answer. Try structuring better using STAR method."
    }

# ── Report generation ──────────────────────────────────────────────────────
def do_generate_report(name: str, role: str, answers: List[Dict]) -> str:
    answered = [a for a in answers if a.get("transcript") not in ("Skipped", "No answer provided")]
    skipped_count = len(answers) - len(answered)
    if not answered:
        return (
            "1. Executive Summary\n"
            f"{name} did not provide enough answered questions for a technical evaluation.\n\n"
            "2. Technical Competency Assessment\nScore: 0/10\nNo technical evidence was captured.\n\n"
            "3. Areas for Development\nAttempt each question with a structured explanation covering approach, implementation, and validation.\n\n"
            "4. Hiring Recommendation\nNo. Insufficient answered evidence."
        )

    def avg(path: str) -> float:
        return round(sum(float(a.get("eval", {}).get(path, 0) or 0) for a in answered) / max(len(answered), 1), 1)

    avg_tech = avg("technical_depth")
    avg_rel = avg("relevance")
    avg_clarity = avg("clarity")
    avg_conf = avg("confidence")
    avg_overall = round(sum(float(a.get("composite", a.get("eval", {}).get("overall", 0)) or 0) for a in answered) / max(len(answered), 1), 1)
    all_strengths = []
    all_improvements = []
    for a in answered:
        all_strengths.extend(a.get("eval", {}).get("strengths", [])[:2])
        all_improvements.extend(a.get("eval", {}).get("improvements", [])[:2])
    strengths = list(dict.fromkeys(all_strengths))[:5] or ["Provided some relevant technical responses."]
    improvements = list(dict.fromkeys(all_improvements))[:5] or ["Add measurable outcomes and deeper implementation detail."]
    recommendation = "Recommend" if avg_overall >= 7.5 else "Maybe" if avg_overall >= 5.5 else "No"

    return (
        "1. Executive Summary\n"
        f"{name} completed {len(answered)} of {len(answers)} technical questions for the {role} role. "
        f"The local NLP evaluator produced an average composite score of {avg_overall}/10. "
        f"{skipped_count} question(s) were skipped or unanswered.\n\n"
        "2. Technical Competency Assessment\n"
        f"Score: {avg_tech}/10. Technical scoring is based on implementation vocabulary, keyword alignment, "
        "semantic similarity to the question, concrete actions, and trade-off discussion.\n\n"
        "3. Communication & Structure\n"
        f"Clarity: {avg_clarity}/10. Confidence: {avg_conf}/10. Relevance: {avg_rel}/10. "
        "Structure scoring rewards step-by-step explanation, examples, specificity, and validation detail.\n\n"
        "4. Key Strengths\n"
        + "\n".join(f"- {s}" for s in strengths)
        + "\n\n5. Areas for Development\n"
        + "\n".join(f"- {i}" for i in improvements)
        + "\n\n6. Hiring Recommendation\n"
        f"{recommendation}. This recommendation is generated from local scoring only, without GPT-based evaluation."
    )

# ── Composite score (synopsis §5 formula) ─────────────────────────────────
def compute_composite(eval_dict, fluency):
    tech = eval_dict.get("technical_depth", 0)
    clarity = eval_dict.get("clarity", 0)
    confidence = eval_dict.get("confidence", 0)
    relevance = eval_dict.get("relevance", 0)
    behavioral = eval_dict.get("behavioral", 0)
    fluency_score = fluency.get("fluency_score", 0)

    # Weighted formula
    composite = (
        tech * 0.4 +
        ((clarity + confidence) / 2) * 0.2 +
        relevance * 0.15 +
        behavioral * 0.15 +
        fluency_score * 0.1
    )

    return round(composite, 1)


# ── Helpers ────────────────────────────────────────────────────────────────
def detect_fmt(data: bytes) -> str:
    if data[:4] == b'RIFF': return "wav"
    if data[:4] == b'OggS': return "ogg"
    if data[:4] == b'\x1aE\xdf\xa3': return "webm"
    return "webm"

def score_color(s) -> str:
    s = float(s)
    if s >= 8: return "#10B981"
    if s >= 6: return "#0F766E"
    if s >= 4: return "#F59E0B"
    return "#EF4444"

def sentiment_pill(s: str) -> str:
    m = {"positive": ("#D1FAE5", "#047857", "#6EE7B7"),
         "neutral":  ("#FEF3C7", "#B45309", "#FDE68A"),
         "negative": ("#FEE2E2", "#B91C1C", "#FCA5A5")}
    bg, fg, br = m.get(s, m["neutral"])
    return f'<span style="background:{bg};color:{fg};border:1px solid {br};padding:3px 10px;border-radius:8px;font-size:.72rem;font-weight:700;">{s.capitalize()}</span>'

def extract_resume_text(resume) -> str:
    text = ""
    debug = []
    try:
        raw = resume.getvalue()
        file_type = (getattr(resume, "type", "") or "").lower()
        file_name = (getattr(resume, "name", "") or "").lower()
        debug.append(f"python={sys.executable}")
        debug.append(f"file={file_name or 'unknown'} type={file_type or 'unknown'} size={len(raw)} bytes")

        if "pdf" in file_type or file_name.endswith(".pdf"):
            try:
                from PyPDF2 import PdfReader
                reader = PdfReader(io.BytesIO(raw))
                for pg in reader.pages:
                    text += (pg.extract_text() or "") + "\n"
                debug.append(f"PyPDF2 extracted {len(text.strip())} characters")
            except Exception as e:
                msg = f"PyPDF2 failed: {e}"
                debug.append(msg)
                log.warning(msg)

            if len(text.strip()) < 40:
                try:
                    import fitz  # PyMuPDF
                    with fitz.open(stream=raw, filetype="pdf") as doc:
                        text = "\n".join(page.get_text("text") for page in doc)
                    debug.append(f"PyMuPDF extracted {len(text.strip())} characters")
                except Exception as e:
                    msg = f"PyMuPDF failed: {e}"
                    debug.append(msg)
                    log.warning(msg)

            if len(text.strip()) < 40:
                try:
                    import shutil
                    import fitz  # PyMuPDF
                    from PIL import Image
                    import pytesseract

                    ocr_pages = []
                    tesseract_cmd = shutil.which("tesseract") or "/opt/homebrew/bin/tesseract"
                    if tesseract_cmd:
                        pytesseract.pytesseract.tesseract_cmd = tesseract_cmd
                    with fitz.open(stream=raw, filetype="pdf") as doc:
                        page_count = min(len(doc), 5)
                        debug.append(f"OCR scanning {page_count} page(s) with {pytesseract.pytesseract.tesseract_cmd}")
                        for page in [doc[i] for i in range(page_count)]:
                            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
                            img = Image.open(io.BytesIO(pix.tobytes("png")))
                            ocr_pages.append(pytesseract.image_to_string(img))
                    text = "\n".join(ocr_pages)
                    debug.append(f"OCR extracted {len(text.strip())} characters")
                except Exception as e:
                    msg = f"OCR failed: {e}"
                    debug.append(msg)
                    log.warning(msg)
        else:
            from docx import Document
            doc = Document(io.BytesIO(raw))
            for p in doc.paragraphs:
                text += p.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    text += " ".join(cell.text for cell in row.cells) + "\n"
            debug.append(f"DOCX extracted {len(text.strip())} characters")
    except Exception as e:
        msg = f"Resume text extraction failed: {e}"
        debug.append(msg)
        log.warning(msg)
    try:
        st.session_state.resume_extract_debug = debug
    except Exception:
        pass
    return text.strip()

def _clean_resume_lines(text: str) -> List[str]:
    """Normalize resume text into non-empty readable lines."""
    text = re.sub(r"\r", "\n", text or "")
    text = re.sub(r"[ \t]+", " ", text)
    lines = []
    for line in text.splitlines():
        line = re.sub(r"\s+", " ", line).strip(" -|•\t")
        if line:
            lines.append(line)
    return lines

def _split_resume_sections(lines: List[str]) -> Dict[str, List[str]]:
    sections = {key: [] for key in RESUME_SECTION_ALIASES}
    current = "summary"
    heading_to_key = {
        alias: key
        for key, aliases in RESUME_SECTION_ALIASES.items()
        for alias in aliases
    }
    for line in lines:
        normalized = re.sub(r"[^a-zA-Z ]", "", line).strip().lower()
        if normalized in heading_to_key and len(line.split()) <= 4:
            current = heading_to_key[normalized]
            continue
        sections.setdefault(current, []).append(line)
    return sections

def _extract_contact(lines: List[str], text: str) -> Dict[str, str]:
    email_match = re.search(r"[\w.+-]+@[\w-]+(?:\.[\w-]+)+", text)
    phone = ""
    for candidate in re.findall(r"\+?\d[\d\s().-]{8,}\d", text):
        digits = re.sub(r"\D", "", candidate)
        if 10 <= len(digits) <= 13:
            phone = candidate.strip(" |,")
            break
    links = re.findall(r"(?:https?://)?(?:www\.)?(?:linkedin\.com/in/[\w-]+|github\.com/[\w-]+|[\w-]+\.github\.io[\w/.-]*)", text, flags=re.I)

    name = ""
    for line in lines[:8]:
        if email_match and email_match.group(0) in line:
            continue
        if phone and phone in line:
            continue
        if "linkedin" in line.lower() or "github" in line.lower():
            continue
        words = line.split()
        if 1 < len(words) <= 5 and not any(ch.isdigit() for ch in line):
            name = line
            break

    return {
        "name": name,
        "email": email_match.group(0) if email_match else "",
        "phone": phone,
        "links": ", ".join(dict.fromkeys(links)),
    }

def _extract_skills(text: str, sections: Dict[str, List[str]]) -> List[str]:
    lower_text = text.lower()
    found = []
    for skill in SKILL_KEYWORDS:
        pattern = r"(?<![a-z0-9+#.])" + re.escape(skill.lower()) + r"(?![a-z0-9+#.])"
        if re.search(pattern, lower_text):
            found.append("scikit-learn" if skill == "sklearn" else skill)

    for line in sections.get("skills", []):
        for part in re.split(r"[,|;/]", line):
            token = part.strip(" .:-").lower()
            if 1 < len(token) <= 30 and not token.startswith(("skills", "technical")):
                found.append(token)

    display_names = {
        "fastapi": "FastAPI", "openai": "OpenAI", "node": "Node.js",
        "javascript": "JavaScript", "typescript": "TypeScript",
        "scikit-learn": "scikit-learn", "pytorch": "PyTorch",
        "tensorflow": "TensorFlow", "opencv": "OpenCV", "power bi": "Power BI",
        "mongodb": "MongoDB", "mysql": "MySQL", "postgresql": "PostgreSQL",
        "sqlite": "SQLite", "github": "GitHub", "rest api": "REST API",
    }
    prettified = []
    for skill in dict.fromkeys(found):
        if skill in display_names:
            prettified.append(display_names[skill])
            continue
        if skill in {"api", "sql", "llm", "rag", "nlp", "aws", "gcp"}:
            prettified.append(skill.upper())
        elif skill in {"html", "css"}:
            prettified.append(skill.upper())
        else:
            prettified.append(skill.title().replace("Ai", "AI").replace("Api", "API"))
    return prettified[:18]

def _top_section_items(lines: List[str], limit: int = 4) -> List[str]:
    items = []
    for line in lines:
        if len(line) < 8:
            continue
        if re.fullmatch(r"[\W_]+", line):
            continue
        items.append(line[:180])
        if len(items) >= limit:
            break
    return items

def parse_resume(text: str) -> Dict:
    """Parse extracted resume text into recruiter-friendly structured fields."""
    lines = _clean_resume_lines(text)
    sections = _split_resume_sections(lines)
    profile = {
        "contact": _extract_contact(lines, text),
        "skills": _extract_skills(text, sections),
        "education": _top_section_items(sections.get("education", []), 3),
        "experience": _top_section_items(sections.get("experience", []), 4),
        "projects": _top_section_items(sections.get("projects", []), 4),
        "certifications": _top_section_items(sections.get("certifications", []), 3),
        "word_count": len(re.findall(r"\b\w+\b", text)),
    }
    if not profile["projects"]:
        project_like = [line for line in lines if re.search(r"\b(project|developed|built|created|implemented)\b", line, re.I)]
        profile["projects"] = _top_section_items(project_like, 4)
    return profile

def resume_context(profile: Dict, text: str) -> str:
    """Create a compact context block for question generation."""
    contact = profile.get("contact", {})
    chunks = [
        f"Candidate name: {contact.get('name') or 'Unknown'}",
        f"Skills: {', '.join(profile.get('skills', [])[:12]) or 'Not clearly listed'}",
        "Education: " + " | ".join(profile.get("education", [])[:2]),
        "Experience: " + " | ".join(profile.get("experience", [])[:3]),
        "Projects: " + " | ".join(profile.get("projects", [])[:3]),
        "\nRaw resume excerpt:\n" + text[:2500],
    ]
    return "\n".join(c for c in chunks if c and not c.endswith(": "))

def _chips(items: List[str], empty: str) -> str:
    if not items:
        return f'<span class="pill p-dim">{html.escape(empty)}</span>'
    return "".join(f'<span class="pill p-v" style="margin:0 6px 6px 0;">{html.escape(item)}</span>' for item in items)

def render_resume_profile(profile: Dict):
    contact = profile.get("contact", {})
    name = html.escape(contact.get("name") or "Name not detected")
    email = html.escape(contact.get("email") or "Email not detected")
    phone = html.escape(contact.get("phone") or "Phone not detected")
    links = html.escape(contact.get("links") or "Links not detected")

    st.markdown(f"""
    <div class="card" style="margin-bottom:1.3rem;">
      <div style="display:flex;justify-content:space-between;gap:1rem;align-items:flex-start;flex-wrap:wrap;">
        <div>
          <div style="font-size:.72rem;font-weight:800;letter-spacing:1px;text-transform:uppercase;color:var(--t3);">Parsed Resume</div>
          <div style="font-size:1.25rem;font-weight:800;font-family:var(--heading);color:var(--t1);margin-top:.2rem;">{name}</div>
        </div>
        <span class="pill p-g">{profile.get("word_count", 0)} words scanned</span>
      </div>
      <div style="margin-top:.8rem;display:grid;grid-template-columns:repeat(auto-fit,minmax(180px,1fr));gap:.55rem;">
        <div class="kv-row"><span class="kv-key">Email</span><span>{email}</span></div>
        <div class="kv-row"><span class="kv-key">Phone</span><span>{phone}</span></div>
        <div class="kv-row"><span class="kv-key">Links</span><span>{links}</span></div>
      </div>
      <div style="margin-top:1rem;">
        <div class="kv-key" style="margin-bottom:.45rem;">Skills detected</div>
        {_chips(profile.get("skills", [])[:14], "No skills detected")}
      </div>
    </div>""", unsafe_allow_html=True)

    cols = st.columns(3)
    section_data = [
        ("Education", profile.get("education", []), "No education section found"),
        ("Experience", profile.get("experience", []), "No experience section found"),
        ("Projects", profile.get("projects", []), "No projects found"),
    ]
    for col, (title, items, empty) in zip(cols, section_data):
        with col:
            body = "".join(f"<li>{html.escape(item)}</li>" for item in items) or f"<li>{html.escape(empty)}</li>"
            st.markdown(f"""
            <div class="card" style="height:100%;padding:1rem 1.1rem;">
              <div style="font-weight:800;font-family:var(--heading);margin-bottom:.5rem;">{title}</div>
              <ul style="padding-left:1.1rem;color:var(--t2);font-size:.86rem;line-height:1.5;">{body}</ul>
            </div>""", unsafe_allow_html=True)

# ── Camera component with MediaPipe Facial Emotion (synopsis §5.2) ─────────
def camera_component():
    """
    Live camera feed with face detection bounding box.
    Uses tracking.js face detection (browser-side, ~200KB, no WASM).
    Draws a green bounding box around detected face — aligned with
    the project report's face detection module (Ch.3, §5 Real-Time Video Streaming).
    """
    cam_html = """
    <style>
      *{box-sizing:border-box;margin:0;padding:0;}
      html,body{background:transparent;font-family:'DM Sans',sans-serif;}
      #wrap{position:relative;width:100%;border-radius:8px;overflow:hidden;background:#111827;}
      #cam-video{width:100%;height:210px;display:block;object-fit:cover;}
      #cam-canvas{position:absolute;top:0;left:0;width:100%;height:210px;pointer-events:none;}
      #face-badge{
        position:absolute;bottom:8px;left:50%;transform:translateX(-50%);
        background:rgba(17,24,39,0.86);color:#fff;font-size:.72rem;font-weight:700;
        padding:4px 12px;border-radius:8px;white-space:nowrap;
        border:1px solid rgba(255,255,255,.15);
      }
      #cam-bar{display:flex;align-items:center;justify-content:space-between;margin-top:6px;
               padding:4px 10px;background:#F8FAFC;border:1px solid #E2E8F0;border-radius:8px;}
      #cam-dot{width:8px;height:8px;border-radius:50%;background:#CBD5E1;flex-shrink:0;}
      #cam-dot.live{background:#10B981;box-shadow:0 0 0 3px rgba(16,185,129,.2);}
      #cam-label{font-size:.72rem;font-weight:700;color:#64748B;}
    </style>
    <div id="wrap">
      <video id="cam-video" autoplay playsinline muted></video>
      <canvas id="cam-canvas"></canvas>
      <div id="face-badge">Starting</div>
    </div>
    <div id="cam-bar">
      <div id="cam-dot"></div>
      <div id="cam-label">Connecting camera</div>
      <div></div>
    </div>
    <!-- tracking.js: lightweight browser face detection, no WASM (~200KB) -->
    <script src="https://cdnjs.cloudflare.com/ajax/libs/tracking.js/1.1.3/tracking-min.js"></script>
    <script src="https://cdnjs.cloudflare.com/ajax/libs/tracking.js/1.1.3/data/face-min.js"></script>
    <script>
    (async function(){
      const video  = document.getElementById('cam-video');
      const canvas = document.getElementById('cam-canvas');
      const ctx    = canvas.getContext('2d');
      const badge  = document.getElementById('face-badge');
      const dot    = document.getElementById('cam-dot');
      const lbl    = document.getElementById('cam-label');

      // Start camera
      try {
        const stream = await navigator.mediaDevices.getUserMedia(
          { video:{ width:320, height:240, frameRate:{ ideal:15, max:20 } }, audio:false }
        );
        video.srcObject = stream;
        dot.classList.add('live');
        lbl.textContent = 'Live face detection';
        lbl.style.color = '#059669';
      } catch(e){
        lbl.textContent = 'Camera unavailable';
        lbl.style.color = '#EF4444';
        badge.textContent = 'No camera';
        return;
      }

      // Face detection loop using canvas snapshot every 1.5s
      // (tracking.js Haar cascade — matches report §5 face detection)
      let tracker;
      try {
        tracker = new tracking.ObjectTracker('face');
        tracker.setInitialScale(4);
        tracker.setStepSize(2);
        tracker.setEdgesDensity(0.1);

        let lastRects = [];
        tracker.on('track', function(e){
          lastRects = e.data;
        });

        setInterval(() => {
          if (video.readyState < 2) return;
          canvas.width  = video.videoWidth  || 320;
          canvas.height = video.videoHeight || 240;
          ctx.clearRect(0, 0, canvas.width, canvas.height);

          // Draw face boxes from last detection
          if (lastRects.length > 0) {
            ctx.strokeStyle = '#10B981';
            ctx.lineWidth   = 2;
            ctx.fillStyle   = 'rgba(16,185,129,0.08)';
            lastRects.forEach(r => {
              ctx.strokeRect(r.x, r.y, r.width, r.height);
              ctx.fillRect(r.x, r.y, r.width, r.height);
            });
            badge.textContent      = 'Face detected';
            badge.style.background = '#10B981cc';
          } else {
            badge.textContent      = 'No face detected';
            badge.style.background = 'rgba(17,24,39,0.86)';
          }

          // Run tracker on current frame
          try {
            const tmpCanvas = document.createElement('canvas');
            tmpCanvas.width  = canvas.width;
            tmpCanvas.height = canvas.height;
            tmpCanvas.getContext('2d').drawImage(video, 0, 0);
            tracking.track(tmpCanvas, tracker);
          } catch(e2){}
        }, 1500);

      } catch(err){
        console.warn('Face tracking init failed:', err);
        badge.textContent = 'Live proctoring';
      }
    })();
    </script>
    """
    st.components.v1.html(cam_html, height=278, scrolling=False)

# ── Session ────────────────────────────────────────────────────────────────
def init():
    defs = {
        "page": "landing", "name": "", "email": "", "role": "Software Engineer",
        "question_pool": [],   # list of tagged question dicts for this session
        "questions": [],       # ordered list of dicts for interview (adaptive)
        "q_idx": 0, "answers": [],
        "spoken": set(), "pending_audio": None, "pending_fmt": "wav",
        "report_text": "", "cam_granted": False,
        "resume_processed": False, "n_personalised": 0,
        "resume_profile": {},
        "elo": 1000.0,         # start easy, climb through medium, reach hard only with strong answers
        "used_q_indices": set(),
        "transcription_failed": False,
    }
    for k, v in defs.items():
        if k not in st.session_state: st.session_state[k] = v
    # Ensure SQLite DB exists on first run
    db_init()

def go(p: str):
    st.session_state.page = p; st.rerun()

# ── Nav ────────────────────────────────────────────────────────────────────
def render_nav():
    nd = st.session_state.name or ""
    user_html = f'<span class="user-chip">{nd}</span>' if nd else ""
    st.markdown(
        f'<div class="top-nav">'
        f'<div class="nav-logo"><div class="nav-gem">NG</div>'
        f'<span class="nav-name">NextGen Recruiter</span></div>'
        f'<div class="nav-right">{user_html}</div>'
        f'</div>', unsafe_allow_html=True
    )

# ── TTS ────────────────────────────────────────────────────────────────────
def speak(text: str):
    import hashlib
    safe = text.replace('"', '\\"').replace('\n', ' ').replace("'", "\\'")
    uid  = hashlib.md5(text.encode()).hexdigest()[:8]
    st.components.v1.html(f"""<script>
    window.speechSynthesis.cancel();
    setTimeout(()=>{{
      window.speechSynthesis.cancel();
      const u=new SpeechSynthesisUtterance("{safe}");
      u.rate=0.84;
      u.pitch=0.92;
      u.volume=1;
      u.lang='en-US';
      function sp(){{
        const vs=window.speechSynthesis.getVoices();
        if(vs.length){{
          const preferred = [
            'Samantha', 'Karen', 'Moira', 'Microsoft Aria',
            'Microsoft Jenny', 'Google US English', 'Google UK English Female',
            'Daniel', 'Alex'
          ];
          u.voice =
            preferred.map(name => vs.find(v => v.name.includes(name))).find(Boolean) ||
            vs.find(v => v.lang === 'en-US' && /female|aria|jenny|samantha|karen/i.test(v.name)) ||
            vs.find(v => v.lang === 'en-GB') ||
            vs.find(v => v.lang === 'en-US') ||
            vs.find(v => v.lang.startsWith('en')) ||
            vs[0];
        }}
        window.speechSynthesis.speak(u);
      }}
      if(window.speechSynthesis.getVoices().length)sp();
      else window.speechSynthesis.onvoiceschanged=sp;
    }},300); /* uid={uid} */
    </script>""", height=0, width=0)

# ══════════════════════════════════════════════════════════════════════════════
# PAGE: LANDING
# ══════════════════════════════════════════════════════════════════════════════
def page_landing():
    render_nav()
    c1, c2 = st.columns([1.15, 1], gap="large")
    with c1:
        st.markdown("""
        <div class="landing-hero">
          <div class="landing-title">Technical screening<br><span>with recruiter-grade focus.</span></div>
        </div>""", unsafe_allow_html=True)
    with c2:
        name  = st.text_input("Full Name *",    value=st.session_state.name,  placeholder="e.g. Shreya Khurana",  key="l_name")
        email = st.text_input("Email *",        value=st.session_state.email, placeholder="shreya@gmail.com",   key="l_email")
        role  = st.text_input("Applying For *", value=st.session_state.role,  placeholder="Software Engineer", key="l_role")
        ready = bool(name.strip() and email.strip() and role.strip())
        st.markdown("<div style='height:.5rem'></div>", unsafe_allow_html=True)
        if st.button("Continue", disabled=not ready, key="btn_land"):
            st.session_state.name  = name.strip()
            st.session_state.email = email.strip()
            st.session_state.role  = role.strip()
            go("setup")

# ══════════════════════════════════════════════════════════════════════════════
# PAGE: SETUP
# ══════════════════════════════════════════════════════════════════════════════
def page_setup():
    render_nav()
    st.markdown("""
    <div class="setup-hero">
      <div class="setup-title">Resume Intake</div>
      <div class="setup-sub">Upload a resume to prepare a focused technical interview.</div>
    </div>""", unsafe_allow_html=True)

    if st.session_state.get("resume_processed"):
        n_p = st.session_state.get("n_personalised", 0)
        st.markdown(f"""
        <div style="background:var(--gg);border:1px solid var(--gb);border-radius:var(--r2);padding:1.5rem 2rem;text-align:center;margin-bottom:1.5rem;">
          <div style="font-weight:700;color:#047857;font-size:1rem;">Resume parsed. {n_p} technical questions prepared.</div>
        </div>""", unsafe_allow_html=True)
        if st.session_state.get("resume_profile"):
            render_resume_profile(st.session_state.resume_profile)
        col1, col3 = st.columns(2)
        with col1:
            if st.button("Upload Different Resume", key="btn_reupload"):
                st.session_state.resume_processed = False
                st.session_state.resume_profile = {}
                st.rerun()
        with col3:
            if st.button("Enter Interview Room", key="btn_start"):
                st.session_state.q_idx = 0; st.session_state.answers = []
                st.session_state.spoken = set(); st.session_state.pending_audio = None
                st.session_state.pending_fmt = "wav"; st.session_state.report_text = ""
                st.session_state.cam_granted = False; st.session_state.elo = 1000.0
                st.session_state.used_q_indices = set(); st.session_state.questions = []
                st.session_state.transcription_failed = False
                for k in [k for k in st.session_state if k.startswith("submitted_")]:
                    del st.session_state[k]
                go("camcheck")
        return

    resume = st.file_uploader("Upload your resume (PDF or DOCX)", type=["pdf", "docx"], key="resume_uploader")
    if resume is not None:
        with st.spinner("Parsing resume and building adaptive question pool..."):
            text = extract_resume_text(resume)
            if not text:
                st.error(
                    "Could not read text from this resume yet. I tried PDF text extraction "
                    "and OCR; the details below show where it failed."
                )
                with st.expander("Resume extraction details", expanded=True):
                    for item in st.session_state.get("resume_extract_debug", []):
                        st.write(item)
                    st.caption(
                        "If OCR extracted 0 characters, the PDF may be a protected file, "
                        "a very low-resolution scan, or a photo with text too small to detect."
                    )
                return
            profile = parse_resume(text)
            personalised = do_gen_questions(resume_context(profile, text), st.session_state.role, INTERVIEW_QUESTION_COUNT)
            base_questions = [
                q for q in QUESTION_BANK
                if q["category"] == "technical" and q["difficulty"] in {"medium", "hard"}
            ]
            pool = [dict(q, idx=i) for i, q in enumerate(base_questions)]
            for i, q in enumerate(personalised):
                q["idx"] = len(base_questions) + i
                pool.append(q)
            st.session_state.question_pool    = pool
            st.session_state.resume_processed = True
            st.session_state.n_personalised   = len(personalised)
            st.session_state.resume_profile   = profile
        st.rerun()

# ══════════════════════════════════════════════════════════════════════════════
# PAGE: CAMERA CHECK
# ══════════════════════════════════════════════════════════════════════════════
def page_camcheck():
    render_nav()
    st.markdown("""
    <div style="text-align:center;padding:2rem 0 1.5rem;">
      <div style="font-size:1.8rem;font-weight:800;font-family:var(--heading);color:var(--t1);margin-bottom:.5rem;">Camera & Microphone Check</div>
      <div style="color:var(--t3);font-size:.95rem;max-width:480px;margin:0 auto;line-height:1.6;">
        Both camera and microphone are <strong>required</strong> for this proctored interview.
        You cannot proceed without granting both.
      </div>
    </div>""", unsafe_allow_html=True)

    # Use session state to track whether JS confirmed both permissions
    if "permissions_granted" not in st.session_state:
        st.session_state.permissions_granted = False

    _, col_c, _ = st.columns([1, 2, 1])
    with col_c:
        st.components.v1.html("""
        <style>
          *{box-sizing:border-box;margin:0;padding:0;}
          body{font-family:'DM Sans',sans-serif;background:transparent;}
          #preview{width:100%;border-radius:8px;background:#111827;display:none;
                   margin-bottom:10px;aspect-ratio:4/3;object-fit:cover;}
          #allow-btn{width:100%;padding:14px;border:none;border-radius:8px;color:#fff;
            font-size:.92rem;font-weight:700;cursor:pointer;display:block;margin-bottom:10px;
            background:#111827;
            box-shadow:none;transition:all .2s;}
          #allow-btn:disabled{background:#CBD5E1;cursor:not-allowed;box-shadow:none;}
          #status{font-size:.84rem;font-weight:600;padding:10px 14px;border-radius:8px;
                  background:#F8FAFC;border:1px solid #E2E8F0;display:none;text-align:center;
                  line-height:1.6;}
        </style>
        <video id="preview" autoplay playsinline muted></video>
        <button id="allow-btn" onclick="setup()">Allow Camera &amp; Microphone</button>
        <div id="status"></div>
        <script>
        async function setup(){
          const ab = document.getElementById('allow-btn');
          const st = document.getElementById('status');
          const pv = document.getElementById('preview');
          ab.disabled = true; ab.textContent = 'Requesting access';
          try {
            const s = await navigator.mediaDevices.getUserMedia({video:true, audio:true});
            pv.srcObject = s;
            pv.style.display = 'block';
            ab.style.display = 'none';
            st.style.display = 'block';
            st.textContent = '';
            st.style.background = 'transparent';
            st.style.borderColor = 'transparent';
          } catch(e) {
            ab.disabled = false;
            ab.textContent = 'Allow Camera & Microphone';
            st.style.display = 'block';
            st.style.color = '#DC2626';
            st.style.background = '#FEF2F2';
            st.style.borderColor = '#FCA5A5';
            if (e.name === 'NotAllowedError') {
              st.textContent = 'Permission denied. Allow camera and microphone in your browser settings, then refresh.';
            } else if (e.name === 'NotFoundError') {
              st.textContent = 'Camera or microphone not found. Please connect both devices.';
            } else {
              st.textContent = e.message;
            }
          }
        }
        </script>""", height=480, scrolling=False)

        confirmed = st.checkbox("Camera and microphone are both active. I am ready.", key="cam_confirmed")
        if confirmed:
            if st.button("Start Interview", key="btn_start_interview"):
                st.session_state.cam_granted = True
                st.session_state.permissions_granted = True
                go("interview")

# ══════════════════════════════════════════════════════════════════════════════
# Interview helpers
# ══════════════════════════════════════════════════════════════════════════════
def get_current_question() -> Optional[Dict]:
    """Return current question dict, picking adaptively if not yet set."""
    qs  = st.session_state.questions
    idx = st.session_state.q_idx
    if idx < len(qs):
        return qs[idx]
    # Need to pick next question adaptively
    pool = st.session_state.question_pool
    used = st.session_state.used_q_indices
    # Load questions this candidate has seen in PREVIOUS sessions from SQLite
    email = st.session_state.get("email", "")
    seen  = db_get_seen_hashes(email) if email else set()
    nxt   = pick_next_question(pool, st.session_state.elo, used, seen)
    if nxt is None: return None
    q_idx_in_pool = nxt.get("idx", id(nxt))
    st.session_state.used_q_indices.add(q_idx_in_pool)
    st.session_state.questions.append(nxt)
    return nxt

def save_answer(transcript: str, fluency: Dict):
    qdict = st.session_state.questions[st.session_state.q_idx]
    st.session_state.answers.append({
        "question":   qdict["q"],
        "category":   qdict.get("category", "general"),
        "difficulty": qdict.get("difficulty", "medium"),
        "personalised": qdict.get("personalised", False),
        "transcript": transcript,
        "fluency":    fluency,
        "eval":       {},
    })

def advance(last_score: float = 0.0):
    # Update Elo based on last answer score (synopsis §5.1)
    if last_score > 0:
        st.session_state.elo = elo_update(st.session_state.elo, last_score)
    total_target = INTERVIEW_QUESTION_COUNT
    if len(st.session_state.answers) >= total_target:
        st.session_state.page = "evaluating"
    else:
        st.session_state.q_idx = len(st.session_state.answers)
    st.session_state.pending_audio = None
    st.session_state.pending_fmt   = "wav"
    st.rerun()

def submit_answer():
    raw = st.session_state.pending_audio
    fmt = st.session_state.pending_fmt

    # Always initialize fluency FIRST (prevents UnboundLocalError)
    fluency = {
        "filler_count": 0,
        "word_count": 0,
        "filler_rate": 0.0,
        "fluency_score": 0,
        "fillers_found": [],
        "pitch_mean": 0.0,
        "pitch_std": 0.0,
        "pause_ratio": 0.0,
        "prosody_score": 0,
        "duration_sec": 0.0
    }

    # If no audio → treat as no answer
    if not raw or len(raw) < 100:
        save_answer("No answer provided", fluency)
        advance(0.0)
        return

    # ── TRANSCRIPTION ─────────────────
    with st.spinner("Transcribing your answer..."):
        try:
            transcript, _ = do_transcribe(raw, fmt)
        except Exception as e:
            log.warning(f"Transcription failed: {e}")
            transcript = ""

    transcript = (transcript or "").strip()
    if not transcript:
        transcript = "No answer provided"

    # ── FLUENCY ANALYSIS ──────────────
    try:
        analysed = analyse_fluency(transcript, raw_audio=raw, fmt=fmt)
        if analysed:
            fluency.update(analysed)
    except Exception as e:
        log.warning(f"Fluency analysis failed: {e}")

    # ── FINAL FLUENCY SCORE ───────────
    fluency["fluency_score"] = enhanced_fluency_score(fluency)

    # Save answer
    st.session_state.pending_audio = None
    st.session_state["transcription_failed"] = False
    st.session_state[f"submitted_{st.session_state.q_idx}"] = True

    save_answer(transcript, fluency)

    # Give small score boost only if answered
    advance(5.0 if transcript != "No answer provided" else 0.0)
def skip_answer():
    fluency = {
        "filler_count": 0,
        "word_count": 0,
        "filler_rate": 0.0,
        "fluency_score": 0,
        "fillers_found": [],
        "pitch_mean": 0.0,
        "pitch_std": 0.0,
        "pause_ratio": 0.0,
        "prosody_score": 0,
        "duration_sec": 0.0
    }

    st.session_state[f"submitted_{st.session_state.q_idx}"] = True

    save_answer("Skipped", fluency)

    # No score for skipped answers
    advance(0.0)

def elo_update(current_elo: float, score: float) -> float:
    """
    Update ELO rating based on answer score (0–10 scale)
    """

    # Normalize score to 0–1
    normalized = score / 10.0

    # Expected performance (baseline = 0.5)
    expected = 0.5

    # K-factor (controls sensitivity)
    K = 32  

    # Elo update formula
    new_elo = current_elo + K * (normalized - expected)

    # Clamp between bounds
    return max(800, min(1600, new_elo))


# ══════════════════════════════════════════════════════════════════════════════
# PAGE: INTERVIEW
# ══════════════════════════════════════════════════════════════════════════════
def page_interview():
    render_nav()
    total_target = INTERVIEW_QUESTION_COUNT
    answered_so_far = len(st.session_state.answers)
    if answered_so_far >= total_target:
        go("evaluating"); return

    qdict = get_current_question()
    if qdict is None:
        go("evaluating"); return

    question     = qdict["q"]
    category     = qdict.get("category", "general")
    difficulty   = qdict.get("difficulty", "medium")
    personalised = qdict.get("personalised", False)
    idx          = st.session_state.q_idx
    progress_pct = int((answered_so_far / total_target) * 100)
    diff_color   = DIFFICULTY_COLORS.get(difficulty, "#64748B")
    elo          = st.session_state.elo

    main_col, panel_col = st.columns([1.6, 1], gap="large")

    with panel_col:
        st.markdown('<div class="panel-label">Live Proctoring</div>', unsafe_allow_html=True)
        camera_component()
        st.markdown(f"""
        <div class="card" style="margin-top:.6rem;padding:1rem 1.2rem;">
          <div class="kv-row">
            <span class="kv-key">Question</span>
            <span class="pill p-v">{answered_so_far+1} / {total_target}</span>
          </div>
          <div class="kv-row">
            <span class="kv-key">Answered</span>
            <span class="pill p-g">{answered_so_far}</span>
          </div>
          <div style="margin-top:.6rem;">
            <div class="prog-wrap"><div class="prog-fill" style="width:{progress_pct}%;"></div></div>
          </div>
        </div>""", unsafe_allow_html=True)

    with main_col:
        st.markdown(
            f'<div class="q-header">'
            f'<span class="q-counter">Question {answered_so_far+1:02d} of {total_target:02d}</span>'
            f'<span class="pill p-dim">{answered_so_far} answered</span>'
            f'</div>', unsafe_allow_html=True
        )
        if idx not in st.session_state.spoken:
            speak(question)
            st.session_state.spoken.add(idx)

        st.markdown(f"""
        <div class="q-box">
          <div class="q-text">{question}</div>
          <div class="q-meta">
            <span class="pill p-dim">{CATEGORY_LABELS.get(category, category)}</span>
            <span style="display:inline-flex;align-items:center;gap:4px;padding:5px 12px;border-radius:8px;font-size:.73rem;font-weight:700;background:#F8FAFC;color:{diff_color};border:1px solid {diff_color}33;">{difficulty.capitalize()}</span>
            {"<span class='pill p-v'>Personalised</span>" if personalised else ""}
          </div>
        </div>""", unsafe_allow_html=True)

        # Audio recorder
        audio_val = st.audio_input("Record your answer", label_visibility="visible", key=f"audio_{idx}")
        if audio_val is not None:
            raw = audio_val.getvalue()
            if raw and len(raw) > 100:
                st.session_state.pending_audio = raw
                st.session_state.pending_fmt   = detect_fmt(raw)
                st.session_state["transcription_failed"] = False

        if st.session_state.get("transcription_failed"):
            st.error(
                "**Transcription failed**. Audio was recorded but could not be converted to text. "
                "Check that ffmpeg is installed and your API keys are correct in secrets.toml. "
                "Re-record and try again, or click Skip to move on."
            )

        has_audio = bool(st.session_state.pending_audio)
        st.markdown("<div style='height:.4rem'></div>", unsafe_allow_html=True)
        c1, c2 = st.columns([3, 1])
        with c1:
            if st.button("Submit Answer", key=f"sub_{idx}", disabled=not has_audio):
                submit_answer()
        with c2:
            if st.button("Skip", key=f"skip_{idx}"):
                st.session_state["transcription_failed"] = False
                skip_answer()

# ══════════════════════════════════════════════════════════════════════════════
# PAGE: EVALUATING
# ══════════════════════════════════════════════════════════════════════════════
def page_evaluating():
    render_nav()
    st.markdown("""
    <div class="eval-screen">
      <div style="font-size:1.8rem;font-weight:800;font-family:var(--heading);color:var(--t1);margin-bottom:.5rem;">Evaluating technical responses</div>

    </div>""", unsafe_allow_html=True)
    answers = st.session_state.answers
    if not answers: go("results"); return
    bar = st.progress(0, text="Starting")
    for i, ans in enumerate(answers):
        bar.progress(i / len(answers), text=f"Analysing answer {i+1} of {len(answers)}")
        if not ans.get("eval"):
            fluency = ans.get("fluency", {})
            ans["eval"] = do_evaluate_answer(ans["question"], ans["transcript"], fluency)
            # Compute synopsis composite score
            ans["composite"] = compute_composite(ans["eval"], fluency)
        time.sleep(0.05)
    bar.progress(0.95, text="Generating evaluation report")
    if not st.session_state.report_text:
        st.session_state.report_text = do_generate_report(
            st.session_state.name, st.session_state.role, answers
        )
    bar.progress(1.0, text="Done!")
    time.sleep(0.4); bar.empty()

    # Persist to SQLite (synopsis §4.2 — database storage)
    answered_evals = [a for a in answers if a.get("composite", 0) > 0]
    final_composite = round(sum(a["composite"] for a in answered_evals) / max(len(answered_evals),1), 1)
    record = {
        "name": st.session_state.name, "email": st.session_state.email,
        "role": st.session_state.role, "date": datetime.now().strftime("%d %b %Y %H:%M"),
        "score": final_composite, "answered": len(answered_evals),
        "total": len(answers), "elo_final": st.session_state.elo,
        "answers": answers, "report": st.session_state.report_text,
    }
    db_save(record)
    # Mark all questions seen this session so they won't repeat next time
    email = st.session_state.email
    if email:
        hashes = [q_hash({"q": a["question"]}) for a in answers]
        db_mark_seen(email, hashes)
    go("results")

# ══════════════════════════════════════════════════════════════════════════════
# PAGE: RESULTS
# ══════════════════════════════════════════════════════════════════════════════
def page_results():
    render_nav()
    answers  = st.session_state.answers
    name     = st.session_state.name
    role     = st.session_state.role
    total    = len(answers)
    answered_list = [a for a in answers if a.get("transcript","") not in ("Skipped","No answer provided")]
    answered = len(answered_list)
    n        = max(answered, 1)

    evals    = [a.get("eval",{}) for a in answered_list]
    fluencies = [a.get("fluency",{}) for a in answered_list]

    # Score = average composite of answered Qs × (answered/total) completion factor
    composites = [a.get("composite", 0) for a in answered_list if a.get("composite",0) > 0]
    avg_composite = round(sum(composites)/max(len(composites),1), 1)
    completion_factor = answered / max(total, 1)
    difficulty_weighted = sum(
        a["composite"] * (1.2 if a.get("difficulty") == "hard" else 1.0)
        for a in answered_list
    ) / max(len(answered_list), 1)

    completion_factor = answered / max(total, 1)

    performance_score = round(avg_composite, 2)
    completion_rate = answered / max(total, 1)

    final_score = round(
        0.7 * performance_score +
        0.3 * (completion_rate * 10),
        2
    )

    avg_tech = round(sum(e.get("technical_depth",0) for e in evals)/n, 1)
    avg_comm = round(sum((e.get("clarity",0)+e.get("confidence",0))/2 for e in evals)/n, 1)
    avg_rel  = round(sum(e.get("relevance",0) for e in evals)/n, 1)
    avg_flu  = round(sum(f.get("fluency_score",0) for f in fluencies)/n, 1)
    sentiments = [e.get("sentiment","neutral") for e in evals]
    pos_pct  = round(sentiments.count("positive")/max(len(sentiments),1)*100)

    tech_scores = [a["composite"] for a in answered_list if a["category"] == "technical"]
    behavior_scores = [a["composite"] for a in answered_list if a["category"] == "behavioral"]

    if behavior_scores:
        avg_behavior_score = round(sum(behavior_scores)/len(behavior_scores), 2)
    else:
        avg_behavior_score = "N/A"

    avg_tech_score = round(sum(tech_scores)/max(len(tech_scores),1),2)
    avg_behavior_score = round(sum(behavior_scores)/max(len(behavior_scores),1),2)
    st.markdown(f"""
    <div style="text-align:center;padding:2rem 0 1.5rem;">
      <div style="font-size:2rem;font-weight:800;font-family:var(--heading);color:var(--t1);margin-bottom:.4rem;">Interview Complete</div>
      <div style="font-size:1rem;color:var(--t2);margin-bottom:1.2rem;">{name} | {role}</div>
      <div style="font-size:.75rem;text-transform:uppercase;letter-spacing:1.5px;font-weight:700;color:var(--t3);margin-bottom:.3rem;">Overall Score</div>
      <div style="font-size:4rem;font-weight:800;font-family:var(--heading);color:{score_color(final_score)};line-height:1;">{final_score}<span style="font-size:1.5rem;color:var(--t3);font-weight:600;"> / 10</span></div>
      <div style="font-size:.82rem;color:var(--t3);margin-top:.5rem;">{answered} of {total} answered</div>
    </div>""", unsafe_allow_html=True)

    st.markdown('<div style="font-size:1.3rem;font-weight:800;font-family:var(--heading);margin:1rem 0 .8rem;">Answer Breakdown</div>', unsafe_allow_html=True)
    for i, ans in enumerate(answers):
        ev  = ans.get("eval", {})
        flu = ans.get("fluency", {})
        comp = ans.get("composite", 0)
        cat  = ans.get("category","")
        diff = ans.get("difficulty","")
        with st.expander(f"Q{i+1} | {ans['question'][:75]}{'...' if len(ans['question'])>75 else ''}", expanded=False):
            transcript_text = ans.get('transcript') or 'No answer recorded.'
            st.markdown(f"""
            <div style="font-size:.78rem;font-weight:700;text-transform:uppercase;letter-spacing:1px;color:var(--t3);margin-bottom:.4rem;">Your Answer</div>
            <div style="background:var(--bg2);border-radius:12px;padding:1rem;font-size:.9rem;color:var(--t2);line-height:1.55;">{transcript_text}</div>
            """, unsafe_allow_html=True)

    st.markdown("<div style='height:1.5rem'></div>", unsafe_allow_html=True)
    if st.button("Start New Interview", key="btn_restart"):
        for k in list(st.session_state.keys()): del st.session_state[k]
        go("landing")
    
    st.write(f"Technical Score: {avg_tech_score}/10")
    st.write(f"Behavioral Score: {avg_behavior_score}/10")

# ══════════════════════════════════════════════════════════════════════════════
# PDF Generation
# ══════════════════════════════════════════════════════════════════════════════
def generate_pdf_report(name, role, answers, report_text, final_score,
                         avg_tech, avg_comm, avg_rel, avg_flu, answered, total) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
    from reportlab.lib.enums import TA_CENTER
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=2*cm, bottomMargin=2*cm, leftMargin=2.5*cm, rightMargin=2.5*cm)
    INDIGO=colors.HexColor("#6366F1"); DARK=colors.HexColor("#0F172A"); MID=colors.HexColor("#334155")
    LIGHT=colors.HexColor("#64748B"); GREEN=colors.HexColor("#10B981"); AMBER=colors.HexColor("#F59E0B")
    RED=colors.HexColor("#EF4444"); BG=colors.HexColor("#F8FAFC")
    def sc(v):
        v=float(v)
        if v>=8: return GREEN
        if v>=6: return INDIGO
        if v>=4: return AMBER
        return RED
    title_s = ParagraphStyle("T",  fontSize=22,fontName="Helvetica-Bold",textColor=DARK, spaceAfter=4,alignment=TA_CENTER)
    sub_s   = ParagraphStyle("S",  fontSize=10,fontName="Helvetica",     textColor=LIGHT,spaceAfter=2,alignment=TA_CENTER)
    h2_s    = ParagraphStyle("H2", fontSize=13,fontName="Helvetica-Bold",textColor=DARK, spaceAfter=6,spaceBefore=12)
    body_s  = ParagraphStyle("B",  fontSize=10,fontName="Helvetica",     textColor=MID,  spaceAfter=4,leading=15)
    q_s     = ParagraphStyle("Q",  fontSize=10,fontName="Helvetica-Bold",textColor=DARK, spaceAfter=3)
    ans_s   = ParagraphStyle("A",  fontSize=9, fontName="Helvetica",     textColor=MID,  spaceAfter=2,leading=13)
    fb_s    = ParagraphStyle("F",  fontSize=9, fontName="Helvetica-Oblique",textColor=LIGHT,spaceAfter=4,leading=13)
    story = []
    story += [Spacer(1,.3*cm), Paragraph("NextGen Recruiter",title_s),
              Paragraph("AI-Powered Interview Evaluation Report",sub_s),
              Paragraph(f"{name}  ·  {role}  ·  {datetime.now().strftime('%d %B %Y')}",sub_s),
              Spacer(1,.4*cm), HRFlowable(width="100%",thickness=2,color=INDIGO), Spacer(1,.4*cm)]
    story.append(Paragraph("Score Summary (Synopsis Formula)", h2_s))
    story.append(Paragraph("Composite = Technical×50% + Communication×25% + Soft Skills×15% + Fluency×10%", body_s))
    score_data = [
        ["Composite","Technical","Communication","Relevance","Fluency","Answered"],
        [str(final_score),str(avg_tech),str(avg_comm),str(avg_rel),str(avg_flu),f"{answered}/{total}"],
    ]
    t = Table(score_data, colWidths=[2.6*cm]*6)
    t.setStyle(TableStyle([
        ("BACKGROUND",(0,0),(-1,0),INDIGO),("TEXTCOLOR",(0,0),(-1,0),colors.white),
        ("FONTNAME",(0,0),(-1,0),"Helvetica-Bold"),("FONTSIZE",(0,0),(-1,0),8),
        ("BACKGROUND",(0,1),(-1,1),BG),("FONTNAME",(0,1),(-1,1),"Helvetica-Bold"),("FONTSIZE",(0,1),(-1,1),14),
        ("TEXTCOLOR",(0,1),(0,1),sc(final_score)),("TEXTCOLOR",(1,1),(1,1),sc(avg_tech)),
        ("TEXTCOLOR",(2,1),(2,1),sc(avg_comm)),("TEXTCOLOR",(3,1),(3,1),sc(avg_rel)),
        ("TEXTCOLOR",(4,1),(4,1),sc(avg_flu)),("TEXTCOLOR",(5,1),(5,1),DARK),
        ("ALIGN",(0,0),(-1,-1),"CENTER"),("VALIGN",(0,0),(-1,-1),"MIDDLE"),
        ("BOX",(0,0),(-1,-1),1,colors.HexColor("#E2E8F0")),
        ("INNERGRID",(0,0),(-1,-1),.5,colors.HexColor("#E2E8F0")),
        ("TOPPADDING",(0,0),(-1,-1),8),("BOTTOMPADDING",(0,0),(-1,-1),8),
    ]))
    story.append(t); story.append(Spacer(1,.5*cm))
    if report_text:
        story += [HRFlowable(width="100%",thickness=1,color=colors.HexColor("#E2E8F0")),
                  Paragraph("AI Evaluation Report",h2_s)]
        for line in report_text.split("\n"):
            line=line.strip()
            if not line: story.append(Spacer(1,.12*cm))
            elif line.startswith(("•","-")): story.append(Paragraph(f"&nbsp;&nbsp;&nbsp;{line}",body_s))
            else: story.append(Paragraph(line,body_s))
        story.append(Spacer(1,.3*cm))
    story += [HRFlowable(width="100%",thickness=1,color=colors.HexColor("#E2E8F0")),
              Paragraph("Question-by-Question Breakdown",h2_s)]
    for i,ans in enumerate(answers):
        ev=ans.get("eval",{}); flu=ans.get("fluency",{})
        cat=ans.get("category",""); diff=ans.get("difficulty","")
        story.append(Paragraph(f"Q{i+1} [{cat.upper()} · {diff.upper()}]: {ans['question']}",q_s))
        story.append(Paragraph(f"Answer: {ans['transcript'] or 'No answer.'}",ans_s))
        if flu.get("word_count",0)>0:
            prosody_str = ""
            if flu.get("pitch_mean",0) > 0:
                prosody_str = f" · Pitch:{flu['pitch_mean']}Hz · Pauses:{round(flu.get('pause_ratio',0)*100,1)}% · Prosody:{flu.get('prosody_score',0)}/10"
            story.append(Paragraph(
                f"Fluency: {flu['word_count']} words · {flu['filler_count']} fillers · Score:{flu.get('fluency_score',0)}/10{prosody_str}",
                fb_s))
        if ev:
            row=[f"Composite:{ans.get('composite',0)}/10",f"Technical:{ev.get('technical_depth','-')}/10",
                 f"Clarity:{ev.get('clarity','-')}/10",f"Confidence:{ev.get('confidence','-')}/10",
                 f"Sentiment:{ev.get('sentiment','-')}"]
            t2=Table([row],colWidths=[3.1*cm]*5)
            t2.setStyle(TableStyle([
                ("FONTNAME",(0,0),(-1,-1),"Helvetica"),("FONTSIZE",(0,0),(-1,-1),8),
                ("TEXTCOLOR",(0,0),(-1,-1),LIGHT),("ALIGN",(0,0),(-1,-1),"CENTER"),
                ("BACKGROUND",(0,0),(-1,-1),BG),("BOX",(0,0),(-1,-1),.5,colors.HexColor("#E2E8F0")),
                ("TOPPADDING",(0,0),(-1,-1),4),("BOTTOMPADDING",(0,0),(-1,-1),4),
            ]))
            story.append(t2)
            if ev.get("feedback"):
                story.append(Paragraph(f"Feedback: {ev['feedback']}",fb_s))
        story.append(Spacer(1,.2*cm))
    story += [HRFlowable(width="100%",thickness=1,color=colors.HexColor("#E2E8F0")),
              Spacer(1,.2*cm),
              Paragraph(f"Generated by NextGen Recruiter · {datetime.now().strftime('%d %b %Y %H:%M')}",sub_s)]
    doc.build(story)
    return buf.getvalue()

# ══════════════════════════════════════════════════════════════════════════════
# MAIN
# ══════════════════════════════════════════════════════════════════════════════
def main():
    st.set_page_config(page_title="NextGen Recruiter", layout="wide", initial_sidebar_state="collapsed")
    st.markdown(CSS, unsafe_allow_html=True)
    init()
    p = st.session_state.page
    if   p == "landing":         page_landing()
    elif p == "setup":           page_setup()
    elif p == "camcheck":        page_camcheck()
    elif p == "interview":       page_interview()
    elif p == "evaluating":      page_evaluating()
    elif p == "results":         page_results()
    else: go("landing")

if __name__ == "__main__":
    main()
